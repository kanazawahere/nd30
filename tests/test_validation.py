"""Test cho engine nd30: build_docx dựng đúng, validate_docx bắt đúng lỗi.

nha-van:exempt — mã nguồn test, không phải văn bản gửi ra ngoài.

Chạy:
    uv venv /tmp/nd30-test-venv
    uv pip install --python /tmp/nd30-test-venv/bin/python python-docx pytest pyyaml
    cd .claude/skills/nd30 && /tmp/nd30-test-venv/bin/python -m pytest tests/ -v
"""
from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "scripts"))

from build_docx import (  # noqa: E402
    add_body_paragraph, add_bullet, add_can_cu, add_header_section, add_kinh_gui,
    add_signature_noi_nhan, add_so_vb_and_date_section, add_table, add_title_block,
    add_two_col_block, placeholder, setup_document,
)
from validate_docx import (  # noqa: E402
    FAIL, OK, WARN, check_auto_bullets, check_placeholders, collect_all_text,
    run_checks, validate_doc,
)

FIXTURES = Path(__file__).resolve().parent / "fixtures"


@pytest.fixture(autouse=True)
def _fixtures_dir():
    FIXTURES.mkdir(parents=True, exist_ok=True)


def _statuses(path, profile="administrative", **kw):
    return {label: status for status, label, _ in run_checks(path, profile, **kw)}


def _full_vb(path: Path) -> Path:
    """Một văn bản hành chính ĐẦY ĐỦ, không còn placeholder → phải sạch lỗi nặng."""
    doc = setup_document(profile="administrative")
    add_header_section(doc, co_quan_chu_quan="UBND HUYỆN A",
                       co_quan_ban_hanh="UBND XÃ B")
    add_so_vb_and_date_section(doc, so_vb="09", ky_hieu="TTr-UBND", dia_danh="Xã B",
                               ngay=5, thang=8, nam=2026)
    add_title_block(doc, ten_loai="TỜ TRÌNH", trich_yeu="Về việc thí điểm mô hình thôn thông minh")
    add_kinh_gui(doc, "Ủy ban nhân dân huyện A")
    add_body_paragraph(doc, "Ủy ban nhân dân xã B kính trình như sau:")
    add_bullet(doc, "Phạm vi triển khai: thôn 1.")
    add_signature_noi_nhan(doc, noi_nhan_items=["- Như trên;", "- Lưu: VT, VP."],
                           chuc_vu="CHỦ TỊCH", nguoi_ky="Nguyễn Văn A")
    doc.save(str(path))
    return path


# =====================================================================
# Tầng A — hình thức
# =====================================================================

def test_van_ban_hanh_chinh_day_du_thi_pass():
    path = _full_vb(FIXTURES / "vb_hop_le.docx")
    assert validate_doc(path, profile="administrative") is True


def test_khong_phai_a4_thi_fail():
    path = FIXTURES / "sai_kho_giay.docx"
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Inches(8.5), Inches(11.0)  # khổ Letter
    doc.save(str(path))
    assert validate_doc(path, verbose=False) is False


def test_le_sai_thi_fail():
    path = FIXTURES / "sai_le.docx"
    doc = setup_document(profile="administrative")
    doc.sections[0].right_margin = Cm(3.0)  # NĐ30: lề phải 1.5–2.0cm
    doc.save(str(path))
    assert validate_doc(path, verbose=False) is False


def test_bullet_ky_tu_tron_thi_fail():
    path = FIXTURES / "bullet_tron.docx"
    doc = setup_document(profile="administrative")
    doc.add_paragraph("• Mục không hợp lệ theo NĐ30")
    doc.save(str(path))
    status, msg = check_auto_bullets(Document(str(path)))
    assert status == FAIL and "•" in msg


def test_bullet_TU_DONG_cua_word_thi_fail():
    """Word numbering (w:numPr) không có ký tự • trong text → phải bắt bằng XML."""
    path = FIXTURES / "bullet_tu_dong.docx"
    doc = setup_document(profile="administrative")
    doc.add_paragraph("Mục dùng style List Bullet", style="List Bullet")
    doc.save(str(path))

    reloaded = Document(str(path))
    # fixture cố ý KHÔNG có ký tự • trong text — lỗi chỉ lộ ra qua style/XML
    assert "•" not in collect_all_text(reloaded)
    assert any(p.style.name.startswith("List Bullet") for p in reloaded.paragraphs)

    status, msg = check_auto_bullets(reloaded)
    assert status == FAIL and "tự động" in msg
    assert validate_doc(path, verbose=False) is False


def test_academic_thi_cho_phep_bullet():
    path = FIXTURES / "academic_bullet.docx"
    doc = setup_document(profile="academic")
    doc.add_paragraph("• Mục hợp lệ trong tài liệu học thuật")
    doc.save(str(path))
    assert validate_doc(path, profile="academic", verbose=False) is True


def test_chu_khong_den_thi_fail():
    path = FIXTURES / "chu_do.docx"
    doc = setup_document(profile="administrative")
    r = doc.add_paragraph().add_run("Chữ màu đỏ")
    r.font.color.rgb = RGBColor(255, 0, 0)
    doc.save(str(path))
    assert validate_doc(path, verbose=False) is False


def test_font_khac_times_new_roman_thi_fail():
    path = FIXTURES / "sai_font.docx"
    doc = setup_document(profile="administrative")
    r = doc.add_paragraph().add_run("Chữ Arial")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    doc.save(str(path))
    assert validate_doc(path, verbose=False) is False


def test_bat_loi_font_mau_TRONG_BANG_va_header():
    """Lỗi nằm trong ô bảng và trong header cũng phải bắt được, không chỉ thân bài."""
    path = FIXTURES / "loi_trong_bang_va_header.docx"
    doc = setup_document(profile="administrative")
    tbl = add_table(doc, ["Cột 1"], [["giá trị"]], widths_cm=[8.0])
    r = tbl.cell(1, 0).paragraphs[0].runs[0]
    r.font.color.rgb = RGBColor(0, 0, 255)
    hr = doc.sections[0].header.paragraphs[0].add_run("Header sai font")
    hr.font.name = "Calibri"
    doc.save(str(path))

    st = _statuses(path)
    assert st["A. Font & màu"] == FAIL
    assert validate_doc(path, verbose=False) is False


def test_bang_tran_le_thi_fail():
    path = FIXTURES / "bang_tran_le.docx"
    doc = setup_document(profile="administrative")
    add_table(doc, ["A", "B"], [["1", "2"]], widths_cm=[12.0, 12.0])  # 24cm > 16cm
    doc.save(str(path))
    st = _statuses(path)
    assert st["A. Bảng tràn lề"] == FAIL


# =====================================================================
# Placeholder
# =====================================================================

def test_bat_placeholder_con_sot():
    path = FIXTURES / "con_placeholder.docx"
    doc = setup_document(profile="administrative")
    add_body_paragraph(doc, "Kinh phí thực hiện: " + placeholder("số tiền"))
    add_table(doc, ["Nội dung"], [["???"]], widths_cm=[8.0])
    doc.save(str(path))

    status, msg = check_placeholders(Document(str(path)))
    assert status == FAIL
    assert "CẦN BỔ SUNG" in msg and "2 placeholder" in msg
    assert validate_doc(path, verbose=False) is False


def test_allow_placeholder_ha_xuong_canh_bao():
    """MẪU trong templates/ vốn phải còn placeholder → chỉ được là ⚠, không phải ✗."""
    path = _full_vb(FIXTURES / "mau_co_placeholder.docx")
    doc = Document(str(path))
    add_body_paragraph(doc, "Người ký: " + placeholder("họ tên người ký"))
    doc.save(str(path))

    assert validate_doc(path, verbose=False) is False
    assert validate_doc(path, verbose=False, allow_placeholder=True) is True
    st = _statuses(path, allow_placeholder=True)
    assert st["A. Placeholder"] == WARN


# =====================================================================
# Helper bảng 2 cột ẩn viền
# =====================================================================

def test_add_two_col_block_khong_co_vien_va_dung_be_rong():
    doc = setup_document(profile="administrative")
    tbl = add_two_col_block(doc, ["Trái"], ["Phải"], left_col_cm=7.0, right_col_cm=9.0)

    assert len(tbl.columns) == 2
    assert abs(tbl.columns[0].width.cm - 7.0) < 0.05
    assert abs(tbl.columns[1].width.cm - 9.0) < 0.05

    borders = tbl._tbl.find(qn("w:tblPr")).find(qn("w:tblBorders"))
    assert borders is not None
    vals = [b.get(qn("w:val")) for b in borders]
    assert set(vals) == {"nil"}, f"viền chưa bị ẩn hết: {vals}"

    assert tbl.cell(0, 0).text.strip() == "Trái"
    assert tbl.cell(0, 1).text.strip() == "Phải"


def test_header_section_quoc_hieu_ben_PHAI_co_quan_ben_TRAI():
    doc = setup_document(profile="administrative")
    tbl = add_header_section(doc, co_quan_chu_quan="UBND HUYỆN A",
                             co_quan_ban_hanh="UBND XÃ B")
    left, right = tbl.cell(0, 0).text, tbl.cell(0, 1).text
    assert "UBND XÃ B" in left and "CỘNG HÒA" not in left
    assert "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in right
    assert "Độc lập - Tự do - Hạnh phúc" in right
    # viền phải ẩn (đây là khối layout, không phải bảng dữ liệu)
    borders = tbl._tbl.find(qn("w:tblPr")).find(qn("w:tblBorders"))
    assert {b.get(qn("w:val")) for b in borders} == {"nil"}


def test_khoi_cuoi_noi_nhan_trai_chu_ky_phai_va_khong_bi_che_trang():
    doc = setup_document(profile="administrative")
    add_signature_noi_nhan(doc, noi_nhan_items=["- Như trên;", "- Lưu: VT, VP."],
                           chuc_vu="CHỦ TỊCH", nguoi_ky="Nguyễn Văn A")
    tbl = doc.tables[-1]
    assert "Nơi nhận:" in tbl.cell(0, 0).text
    assert "Lưu: VT, VP." in tbl.cell(0, 0).text  # giữ dấu chấm cuối theo NĐ30
    assert "CHỦ TỊCH" in tbl.cell(0, 1).text and "Nguyễn Văn A" in tbl.cell(0, 1).text
    trPr = tbl.rows[0]._tr.find(qn("w:trPr"))
    assert trPr is not None and trPr.find(qn("w:cantSplit")) is not None


# =====================================================================
# Căn cứ + 9 thành phần
# =====================================================================

def test_add_can_cu_dat_dau_dung():
    doc = setup_document(profile="administrative")
    ps = add_can_cu(doc, ["Căn cứ Luật Tổ chức chính quyền địa phương",
                          "Căn cứ Nghị định 30/2020/NĐ-CP",
                          "Theo đề nghị của Văn phòng"])
    assert ps[0].text.endswith(";")
    assert ps[1].text.endswith(";")
    assert ps[2].text.endswith(".") and not ps[2].text.endswith(";.")
    assert all(r.italic for p in ps for r in p.runs)  # NĐ30: căn cứ in nghiêng


def test_thieu_noi_nhan_thi_fail():
    path = FIXTURES / "thieu_noi_nhan.docx"
    doc = setup_document(profile="administrative")
    add_header_section(doc, co_quan_chu_quan="UBND HUYỆN A", co_quan_ban_hanh="UBND XÃ B")
    add_so_vb_and_date_section(doc, so_vb="09", ky_hieu="TTr-UBND", dia_danh="Xã B",
                               ngay=5, thang=8, nam=2026)
    add_title_block(doc, ten_loai="TỜ TRÌNH", trich_yeu="Về việc thí điểm")
    add_body_paragraph(doc, "Nội dung trình bày.")
    doc.save(str(path))
    st = _statuses(path)
    assert st["B8. Nơi nhận + Lưu"] == FAIL


def test_bieu_mau_noi_bo_duoc_mien_so_vb_va_noi_nhan():
    path = FIXTURES / "phieu_bieu_quyet.docx"
    doc = setup_document(profile="administrative")
    add_header_section(doc, co_quan_chu_quan="UBND TỈNH X", co_quan_ban_hanh="VĂN PHÒNG")
    add_title_block(doc, ten_loai="PHIẾU BIỂU QUYẾT", trich_yeu="Về dự thảo nghị quyết")
    add_body_paragraph(doc, "Ý kiến: đồng ý.")
    doc.save(str(path))
    st = _statuses(path, profile="bieu-mau-noi-bo")
    assert st["B3. Số/ký hiệu"] == OK
    assert st["B8. Nơi nhận + Lưu"] == OK


def test_profile_general_khong_doi_9_thanh_phan():
    path = FIXTURES / "ghi_chu_noi_bo.docx"
    doc = setup_document(profile="general")
    add_body_paragraph(doc, "Ghi chú nội bộ, không phải văn bản hành chính.")
    doc.save(str(path))
    st = _statuses(path, profile="general")
    assert not any(k.startswith("B") for k in st)
    assert validate_doc(path, profile="general", verbose=False) is True


# =====================================================================
# Templates kèm skill
# =====================================================================

@pytest.mark.parametrize("name", ["to-trinh.docx", "cong-van.docx", "quyet-dinh.docx",
                                  "ke-hoach.docx", "bao-cao.docx"])
def test_mau_trong_templates_tu_pass(name):
    path = ROOT / "templates" / name
    assert path.is_file(), f"thiếu mẫu {name} — chạy templates/_build_templates.py"
    assert validate_doc(path, profile="administrative", verbose=False,
                        allow_placeholder=True) is True
