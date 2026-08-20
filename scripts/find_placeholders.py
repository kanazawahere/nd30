"""find_placeholders.py — scan 1 file .docx tìm các placeholders [KEY] để AI biết
cần fill những gì khi dùng template.

Patterns nhận diện:
  - `[UPPER_SNAKE]` — vd `[TEN_CQ_BAN_HANH]`, `[NOI_DUNG_I_1]`
  - `[Tiếng Việt có space]` — vd `[Phòng soạn thảo văn bản này]`
  - `[CAU_KET_DONG_1]`, `[NOI_NHAN_2]` etc.

KHÔNG nhận diện (tránh false positive):
  - `[1]`, `[A]` — quá ngắn, có thể là footnote/reference
  - `[I.II.III]` — có dấu chấm = không phải placeholder

Usage:
    python find_placeholders.py <file.docx>
    python find_placeholders.py <file.docx> --json

Module API:
    from find_placeholders import scan
    result = scan(Path("template.docx"))
    # result = {"placeholders": [...], "total_unique": int}
"""
from __future__ import annotations

import argparse
import json
import re
import sys, io
from pathlib import Path

if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document


# Bracket có ≥3 ký tự bên trong, không có dấu chấm/comma/colon (tránh "[1.5]", "[1,2]")
PLACEHOLDER_RE = re.compile(
    # (1) placeholder chuẩn của skill nd30: [CẦN BỔ SUNG: ...] — có dấu hai chấm bên trong
    r"\[(CẦN BỔ SUNG[^\[\]]{0,120})\]"
    # (2) placeholder kiểu template cũ: [UPPER_SNAKE] / [Tiếng Việt có space] — không chứa . , :
    r"|\[([^\[\]\.,:]{3,})\]"
)


def _scan_text(text: str) -> list[str]:
    if not text:
        return []
    # regex có 2 nhóm (nd30 / template cũ) → lấy nhóm nào khớp
    matches = [a or b for a, b in PLACEHOLDER_RE.findall(text)]
    # Lọc thêm: phải có ít nhất 1 chữ cái (không chỉ số/space)
    return [m for m in matches if re.search(r"[A-Za-zÀ-ỹĐđ]", m)]


def scan(path: Path) -> dict:
    """Scan file .docx, return {placeholders: [...], total_unique: N}."""
    doc = Document(str(path))
    occurrences: dict[str, list[dict]] = {}

    def add(key: str, location: str, snippet: str):
        bracket = f"[{key}]"
        if bracket not in occurrences:
            occurrences[bracket] = []
        # Truncate snippet for readability
        snip = snippet.strip()
        if len(snip) > 80:
            snip = snip[:77] + "..."
        occurrences[bracket].append({"location": location, "snippet": snip})

    # Top-level paragraphs
    for pi, p in enumerate(doc.paragraphs):
        text = "".join(r.text for r in p.runs)
        for m in _scan_text(text):
            add(m, f"paragraph_{pi}", text)

    # Table cells
    for ti, t in enumerate(doc.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    text = "".join(r.text for r in p.runs)
                    for m in _scan_text(text):
                        add(m, f"table_{ti}/r_{ri}/c_{ci}/p_{pi}", text)

    placeholders = []
    for key in sorted(occurrences.keys()):
        placeholders.append({
            "key": key,
            "count": len(occurrences[key]),
            "occurs": occurrences[key],
        })

    return {
        "file": path.name,
        "placeholders": placeholders,
        "total_unique": len(placeholders),
        "total_occurrences": sum(p["count"] for p in placeholders),
    }


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("file")
    ap.add_argument("--json", action="store_true", help="Output JSON thay vì human-readable")
    args = ap.parse_args()

    path = Path(args.file).resolve()
    if not path.is_file():
        print(f"ERROR: file not found: {path}", file=sys.stderr)
        return 1

    result = scan(path)

    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return 0

    # Human-readable
    print(f"=== Placeholders in {result['file']} ===")
    print(f"  Unique: {result['total_unique']} | Total occurrences: {result['total_occurrences']}")
    print()
    if not result["placeholders"]:
        print("  (không tìm thấy placeholder nào theo pattern [KEY])")
        return 0
    for p in result["placeholders"]:
        print(f"  {p['key']}  ({p['count']} lần)")
        for occ in p["occurs"][:3]:
            print(f"      ↳ {occ['location']}: {occ['snippet']}")
        if p["count"] > 3:
            print(f"      ↳ ... +{p['count']-3} more")
    return 0


if __name__ == "__main__":
    sys.exit(main())
