"""Validate 1 file .docx theo thể thức Nghị định 30/2020/NĐ-CP.

nha-van:exempt — đây là mã nguồn, không phải văn bản gửi ra ngoài.

Hai tầng kiểm:
  A. HÌNH THỨC (đo được chắc chắn từ XML): khổ giấy, lề, font, màu chữ, bullet tự động
     của Word, shading ô bảng, bảng tràn lề, placeholder còn sót.
  B. THÀNH PHẦN THỂ THỨC (9 mục, heuristic theo text): quốc hiệu/tiêu ngữ, tên cơ quan,
     số-ký hiệu, tên loại + trích yếu, nội dung, người ký, dấu, nơi nhận, phụ lục.

Usage:
    python3 validate_docx.py <file.docx> [--profile administrative]

Exit code: 0 = sạch · 1 = chỉ có cảnh báo (⚠) · 2 = có lỗi nặng (✗)

Profile: xem references/document-profiles.md
    administrative (default) · bieu-mau-noi-bo · minutes-administrative · academic · general

Nguồn: tầng B port từ repo vbhc (Unlicense) — luật đọc từ scripts/rules/the-thuc.yaml,
thiếu YAML thì dùng fallback hardcode dưới đây. Tầng A tự viết theo NĐ30 Phụ lục I.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.resolve()))

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

from _common import slugify_vn  # noqa: F401,E402  # kéo theo fix stdout UTF-8
from rules_loader import load_rules  # noqa: E402

OK = "✓"
WARN = "⚠"
FAIL = "✗"

# Profile nào áp luật văn bản hành chính (cấm bullet tự động, cấm shading, đòi 9 thành phần)
ADMIN_PROFILES = {"administrative", "bieu-mau-noi-bo", "minutes-administrative"}

# Placeholder pattern của skill này (bổ sung [CẦN BỔ SUNG] so với bản gốc)
PLACEHOLDER_PATTERNS = [
    r"\[CẦN BỔ SUNG[^\]]*\]",
    r"\?\?\?",
    r"<[^>]{2,40}>",
    r"\[placeholder\]",
]


# =====================================================================
# Fallback luật thể thức (dùng khi scripts/rules/the-thuc.yaml thiếu)
# =====================================================================
_FALLBACK = {
    "placeholder_pattern": "|".join(PLACEHOLDER_PATTERNS),
    "bieu_mau_noi_bo_pattern":
        r"PHIẾU\s+(GHI\s+Ý\s+KIẾN|BIỂU\s+QUYẾT|THẨM\s+ĐỊNH|LẤY\s+Ý\s+KIẾN)",
    "checks": {
        "quoc_hieu": {
            "must_contain": ["CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
                             "Độc lập", "Tự do", "Hạnh phúc"],
            "ok_msg": "Quốc hiệu + Tiêu ngữ",
            "fail_msg": "Quốc hiệu + Tiêu ngữ — thiếu hoặc sai",
        },
        "co_quan": {
            "upper_line_keywords_regex":
                r"(UBND|ỦY BAN|BỘ|SỞ|HĐND|VĂN PHÒNG|TỔNG CỤC|CỤC|CHI CỤC|VIỆN|BAN|TRƯỜNG)",
            "min_line_len": 4,
            "ok_msg": "Tên cơ quan ban hành",
            "warn_msg": "Tên cơ quan — không phát hiện được dòng tên cơ quan in hoa",
        },
        "so_van_ban": {
            "skip_if_bieu_mau_noi_bo": True,
            "pattern": r"Số:\s*([0-9]*\s*/[A-ZĐa-z&\-_]+)",
            "has_so_marker": "Số:",
            "bieu_mau_msg": "Số VB — không cần (biểu mẫu nội bộ)",
            "ok_template": "Số văn bản: {match}",
            "empty_warn_template": "Số văn bản đang trống ({match}) — văn thư sẽ điền khi ban hành",
            "no_format_warn": "Có dòng 'Số:' nhưng không đúng format Số: <số>/<KÝ HIỆU>",
            "fail_msg": "Không tìm thấy 'Số:' — kiểm tra lại",
        },
        "ten_loai": {
            "keywords": ["QUYẾT ĐỊNH", "NGHỊ QUYẾT", "BÁO CÁO", "TỜ TRÌNH",
                         "THÔNG BÁO", "KẾ HOẠCH", "CHỈ THỊ", "BIÊN BẢN",
                         "GIẤY MỜI", "PHIẾU GHI Ý KIẾN", "PHIẾU BIỂU QUYẾT",
                         "HƯỚNG DẪN", "KẾT LUẬN", "ĐỀ ÁN", "PHƯƠNG ÁN"],
            "cong_van_markers": ["V/v", "V/V"],
            "ok_template": "Tên loại: {match}",
            "cong_van_msg": "Tên loại: Công văn (có trích yếu V/v...)",
            "warn_msg": "Không xác định được loại văn bản",
        },
        "noi_dung": {
            "ok_msg": "Nội dung — không còn placeholder",
            "fail_template": "Còn placeholder: {matches}",
        },
        "nguoi_ky": {
            "pattern": r"(KT\.|TL\.|TUQ\.|TM\.|Q\.)?\s*(GIÁM ĐỐC|CHỦ TỊCH|PHÓ CHỦ TỊCH|"
                       r"CHÁNH VĂN PHÒNG|PHÓ CHÁNH VĂN PHÒNG|TRƯỞNG PHÒNG|"
                       r"PHÓ GIÁM ĐỐC|BỘ TRƯỞNG|THỨ TRƯỞNG|VỤ TRƯỞNG|"
                       r"CỤC TRƯỞNG|VIỆN TRƯỞNG|HIỆU TRƯỞNG)",
            "ok_template": "Chức vụ người ký: {match}",
            "warn_msg": "Không phát hiện được chức vụ người ký in hoa",
        },
        "dau": {
            "msg": "Dấu / chữ ký số — script không kiểm được, phải soi tay",
            "status": "warn",
        },
        "noi_nhan": {
            "skip_if_bieu_mau_noi_bo": True,
            "noi_nhan_markers": ["Nơi nhận:", "Nơi nhận :"],
            "luu_markers": ["Lưu:", "Lưu :"],
            "bieu_mau_msg": "Nơi nhận — không cần (biểu mẫu nội bộ)",
            "ok_msg": "Nơi nhận + Lưu",
            "missing_luu_warn": "Có 'Nơi nhận' nhưng thiếu dòng 'Lưu: VT, ...'",
            "fail_msg": "Thiếu 'Nơi nhận:'",
        },
        "phu_luc": {
            "kem_pattern": r"kèm theo|đính kèm",
            "phu_luc_pattern": r"PHỤ LỤC\s+[IVX]+",
            "has_kem_no_pl_warn": "Văn bản nhắc 'kèm theo' nhưng không thấy PHỤ LỤC trong file",
            "has_pl_ok": "Có Phụ lục",
            "no_kem_no_pl_ok": "Không có phụ lục (không bắt buộc)",
        },
    },
}


def _rules() -> dict:
    data = load_rules("the-thuc")
    return data if data else _FALLBACK


def _check_cfg(name: str) -> dict:
    return _rules().get("checks", {}).get(name) or _FALLBACK["checks"][name]


# =====================================================================
# Helper đọc text
# =====================================================================

def collect_all_text(doc) -> str:
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    for section in doc.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
    return "\n".join(parts)


def placeholder_re() -> re.Pattern:
    pat = _rules().get("placeholder_pattern") or _FALLBACK["placeholder_pattern"]
    # Bản YAML gốc không biết [CẦN BỔ SUNG] → luôn ghép thêm pattern của skill này
    if "CẦN BỔ SUNG" not in pat:
        pat = PLACEHOLDER_PATTERNS[0] + "|" + pat
    return re.compile(pat, re.IGNORECASE)


def is_bieu_mau_noi_bo(text: str) -> bool:
    pat = _rules().get("bieu_mau_noi_bo_pattern") or _FALLBACK["bieu_mau_noi_bo_pattern"]
    return bool(re.search(pat, text))


def _iter_all_paragraphs(doc):
    """Mọi paragraph: thân bài + trong ô bảng (kể cả bảng lồng) + header/footer."""
    for i, p in enumerate(doc.paragraphs):
        yield f"P{i}", p
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    yield f"T{ti}/R{ri}/C{ci}/P{pi}", p
                for sti, sub in enumerate(cell.tables):
                    for sri, srow in enumerate(sub.rows):
                        for sci, scell in enumerate(srow.cells):
                            for spi, sp in enumerate(scell.paragraphs):
                                yield f"T{ti}/R{ri}/C{ci}/T{sti}/R{sri}/C{sci}/P{spi}", sp
    for si, section in enumerate(doc.sections):
        for label, container in (("header", section.header), ("footer", section.footer)):
            for pi, p in enumerate(container.paragraphs):
                yield f"S{si}/{label}/P{pi}", p


# =====================================================================
# Tầng A — hình thức đo được
# =====================================================================

def check_page_geometry(doc) -> list[tuple[str, str]]:
    """Khổ giấy + lề theo NĐ30 PL I mục I.1 và I.3."""
    out: list[tuple[str, str]] = []
    if not doc.sections:
        return [(WARN, "Không đọc được section nào")]
    for i, s in enumerate(doc.sections):
        w, h = s.page_width.cm, s.page_height.cm
        if abs(w - 21.0) > 0.5 or abs(h - 29.7) > 0.5:
            out.append((FAIL, f"Section {i}: khổ giấy không phải A4 ({w:.2f}×{h:.2f} cm)"))
        else:
            out.append((OK, f"Section {i}: khổ A4 ({w:.2f}×{h:.2f} cm)"))
        bad = []
        for label, val, lo, hi in (
            ("lề trái", s.left_margin.cm, 3.0, 3.5),
            ("lề phải", s.right_margin.cm, 1.5, 2.0),
            ("lề trên", s.top_margin.cm, 2.0, 2.5),
            ("lề dưới", s.bottom_margin.cm, 2.0, 2.5),
        ):
            if not (lo - 0.05 <= val <= hi + 0.05):
                bad.append(f"{label} {val:.2f}cm (cần {lo}–{hi}cm)")
        if bad:
            out.append((FAIL, f"Section {i}: lề sai NĐ30 — " + "; ".join(bad)))
        else:
            out.append((OK, f"Section {i}: lề đúng dải NĐ30 "
                            f"(T{s.left_margin.cm:.1f}/P{s.right_margin.cm:.1f}/"
                            f"Tr{s.top_margin.cm:.1f}/D{s.bottom_margin.cm:.1f} cm)"))
    return out


def check_font_and_color(doc) -> list[tuple[str, str]]:
    """Font Times New Roman + màu đen ở MỌI run (thân bài, bảng, header/footer)."""
    bad_font, bad_color = [], []
    for loc, p in _iter_all_paragraphs(doc):
        for r in p.runs:
            if r.font.name and r.font.name != "Times New Roman":
                bad_font.append(f"{loc}:{r.font.name}")
            color = r.font.color
            if color is not None and color.rgb is not None and str(color.rgb) != "000000":
                bad_color.append(f"{loc}:#{color.rgb}")
    out = []
    if bad_font:
        out.append((FAIL, f"Font khác Times New Roman ({len(bad_font)} chỗ): {bad_font[:4]}"))
    else:
        out.append((OK, "Font Times New Roman toàn văn bản"))
    if bad_color:
        out.append((FAIL, f"Chữ không đen ({len(bad_color)} chỗ): {bad_color[:4]}"))
    else:
        out.append((OK, "Màu chữ đen toàn văn bản"))
    return out


def check_auto_bullets(doc) -> tuple[str, str]:
    """NĐ30: văn bản hành chính không dùng bullet/numbering tự động của Word."""
    hits = []
    for loc, p in _iter_all_paragraphs(doc):
        if "•" in p.text:
            hits.append(f"{loc}: ký tự •")
            continue
        pPr = p._p.pPr
        if pPr is not None and pPr.find(qn("w:numPr")) is not None:
            hits.append(f"{loc}: numbering tự động của Word")
            continue
        # Numbering có thể nằm trong ĐỊNH NGHĨA STYLE (List Bullet / List Number),
        # lúc đó pPr của paragraph rỗng → phải soi cả style.
        style = p.style
        style_name = getattr(style, "name", "") or ""
        if style_name.startswith(("List Bullet", "List Number", "List Paragraph")):
            hits.append(f"{loc}: style '{style_name}' (bullet/numbering tự động)")
            continue
        st_el = getattr(style, "element", None)
        if st_el is not None:
            st_pPr = st_el.find(qn("w:pPr"))
            if st_pPr is not None and st_pPr.find(qn("w:numPr")) is not None:
                hits.append(f"{loc}: style '{style_name}' có numbering tự động")
    if hits:
        return FAIL, (f"Bullet/numbering tự động ({len(hits)} chỗ): {hits[:4]}"
                      " — dùng tiền tố tay '-', '+', 'a)'")
    return OK, "Không có bullet tự động (dùng tiền tố tay)"


def check_table_shading(doc) -> tuple[str, str]:
    hits = []
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                shd = cell._tc.get_or_add_tcPr().find(f"{ns}shd")
                if shd is None:
                    continue
                val, fill = shd.get(f"{ns}val"), shd.get(f"{ns}fill")
                if (val not in (None, "clear")) or (fill not in (None, "auto", "000000", "FFFFFF")):
                    hits.append(f"T{ti}/R{ri}/C{ci}")
    if hits:
        return FAIL, f"Ô bảng bị tô nền ({len(hits)} ô): {hits[:4]}"
    return OK, "Bảng không tô nền"


def check_table_width(doc) -> tuple[str, str]:
    """Bảng tràn lề: tổng bề rộng cột > vùng nội dung (21cm − lề trái − lề phải)."""
    if not doc.tables:
        return OK, "Không có bảng"
    s = doc.sections[0]
    content_cm = s.page_width.cm - s.left_margin.cm - s.right_margin.cm
    hits = []
    for ti, table in enumerate(doc.tables):
        total = 0.0
        for col in table.columns:
            if col.width is not None:
                total += col.width.cm
        if total and total > content_cm + 0.2:
            hits.append(f"T{ti}: {total:.2f}cm > {content_cm:.2f}cm")
    if hits:
        return FAIL, f"Bảng tràn lề: {hits}"
    return OK, f"Bảng nằm trong vùng nội dung ({content_cm:.2f}cm)"


def check_placeholders(doc) -> tuple[str, str]:
    rx = placeholder_re()
    hits = []
    for loc, p in _iter_all_paragraphs(doc):
        for m in rx.finditer(p.text):
            hits.append(f"{loc}: {m.group(0)}")
    if hits:
        return FAIL, f"Còn {len(hits)} placeholder chưa điền: {hits[:5]}"
    return OK, "Không còn placeholder"


# =====================================================================
# Tầng B — 9 thành phần thể thức
# =====================================================================

def check_quoc_hieu(text):
    cfg = _check_cfg("quoc_hieu")
    if all(s in text for s in cfg["must_contain"]):
        return OK, cfg["ok_msg"]
    return FAIL, cfg["fail_msg"]


def check_co_quan(text):
    cfg = _check_cfg("co_quan")
    min_len = int(cfg.get("min_line_len", 4))
    upper_lines = [ln for ln in text.split("\n")
                   if ln.strip() and ln == ln.upper() and len(ln) > min_len]
    if any(re.search(cfg["upper_line_keywords_regex"], ln) for ln in upper_lines):
        return OK, cfg["ok_msg"]
    return WARN, cfg["warn_msg"]


def check_so_van_ban(text):
    cfg = _check_cfg("so_van_ban")
    if cfg.get("skip_if_bieu_mau_noi_bo") and is_bieu_mau_noi_bo(text):
        return OK, cfg["bieu_mau_msg"]
    m = re.search(cfg["pattern"], text)
    if m:
        if m.group(1).strip().startswith("/"):
            return WARN, cfg["empty_warn_template"].format(match=m.group(0).strip())
        return OK, cfg["ok_template"].format(match=m.group(0).strip())
    if cfg.get("has_so_marker") and cfg["has_so_marker"] in text:
        return WARN, cfg["no_format_warn"]
    return FAIL, cfg["fail_msg"]


def check_ten_loai(text):
    cfg = _check_cfg("ten_loai")
    found = [k for k in cfg["keywords"] if k in text]
    if found:
        return OK, cfg["ok_template"].format(match=found[0])
    if any(m in text for m in cfg.get("cong_van_markers", [])):
        return OK, cfg["cong_van_msg"]
    return WARN, cfg["warn_msg"]


def check_noi_dung(text):
    cfg = _check_cfg("noi_dung")
    rx = placeholder_re()
    if rx.search(text):
        return FAIL, cfg["fail_template"].format(matches=[m.group(0) for m in rx.finditer(text)][:3])
    return OK, cfg["ok_msg"]


def check_nguoi_ky(text):
    cfg = _check_cfg("nguoi_ky")
    m = re.search(cfg["pattern"], text)
    if m:
        return OK, cfg["ok_template"].format(match=m.group(0).strip())
    return WARN, cfg["warn_msg"]


def check_dau():
    cfg = _check_cfg("dau")
    return {"ok": OK, "warn": WARN, "fail": FAIL}.get(cfg.get("status", "warn"), WARN), cfg["msg"]


def check_noi_nhan(text):
    cfg = _check_cfg("noi_nhan")
    if cfg.get("skip_if_bieu_mau_noi_bo") and is_bieu_mau_noi_bo(text):
        return OK, cfg["bieu_mau_msg"]
    has_nn = any(m in text for m in cfg["noi_nhan_markers"])
    has_luu = any(m in text for m in cfg["luu_markers"])
    if has_nn:
        return (OK, cfg["ok_msg"]) if has_luu else (WARN, cfg["missing_luu_warn"])
    return FAIL, cfg["fail_msg"]


def check_phu_luc(text):
    cfg = _check_cfg("phu_luc")
    has_kem = bool(re.search(cfg["kem_pattern"], text, re.IGNORECASE))
    has_pl = bool(re.search(cfg["phu_luc_pattern"], text))
    if has_kem and not has_pl:
        return WARN, cfg["has_kem_no_pl_warn"]
    return (OK, cfg["has_pl_ok"]) if has_pl else (OK, cfg["no_kem_no_pl_ok"])


# =====================================================================
# Điều phối
# =====================================================================

def run_checks(filepath, profile: str = "administrative",
               allow_placeholder: bool = False) -> list[tuple[str, str, str]]:
    """Trả list (status, label, detail).

    allow_placeholder=True → placeholder chỉ là cảnh báo (⚠) thay vì lỗi nặng (✗).
    Dùng cho MẪU trong templates/ (mẫu vốn phải còn placeholder). Bản giao khách thì
    KHÔNG bật cờ này.
    """
    doc = Document(str(filepath))
    text = collect_all_text(doc)
    is_admin = profile in ADMIN_PROFILES
    results: list[tuple[str, str, str]] = []

    # --- Tầng A ---
    for status, msg in check_page_geometry(doc):
        results.append((status, "A. Trang giấy & lề", msg))
    for status, msg in check_font_and_color(doc):
        results.append((status, "A. Font & màu", msg))
    st, msg = check_placeholders(doc)
    if allow_placeholder and st == FAIL:
        st, msg = WARN, msg + "  [chấp nhận: đang validate MẪU]"
    results.append((st, "A. Placeholder", msg))
    st, msg = check_table_width(doc)
    results.append((st, "A. Bảng tràn lề", msg))
    if is_admin:
        st, msg = check_auto_bullets(doc)
        results.append((st, "A. Bullet tự động", msg))
        st, msg = check_table_shading(doc)
        results.append((st, "A. Shading bảng", msg))

    # --- Tầng B: chỉ áp cho văn bản hành chính ---
    if is_admin:
        bieu_mau = profile == "bieu-mau-noi-bo" or is_bieu_mau_noi_bo(text)
        b = [
            ("B1. Quốc hiệu + Tiêu ngữ", check_quoc_hieu(text)),
            ("B2. Tên cơ quan ban hành", check_co_quan(text)),
            ("B3. Số/ký hiệu", (OK, "Số VB — không cần (biểu mẫu nội bộ)") if bieu_mau
             else check_so_van_ban(text)),
            ("B4. Tên loại + trích yếu", check_ten_loai(text)),
            ("B5. Nội dung", (WARN, "Nội dung — mẫu còn placeholder (chấp nhận với templates/)")
             if allow_placeholder and check_noi_dung(text)[0] == FAIL else check_noi_dung(text)),
            ("B6. Người ký", check_nguoi_ky(text)),
            ("B7. Dấu/chữ ký số", check_dau()),
            ("B8. Nơi nhận + Lưu", (OK, "Nơi nhận — không cần (biểu mẫu nội bộ)") if bieu_mau
             else check_noi_nhan(text)),
            ("B9. Phụ lục", check_phu_luc(text)),
        ]
        for label, (st2, detail) in b:
            results.append((st2, label, detail))
    return results


def validate_doc(filepath, profile: str = "administrative", verbose: bool = True,
                 allow_placeholder: bool = False) -> bool:
    """API cho test / skill khác: True nếu KHÔNG có lỗi nặng (✗). Cảnh báo (⚠) vẫn True."""
    try:
        results = run_checks(filepath, profile, allow_placeholder)
    except Exception as e:  # file hỏng / không mở được
        if verbose:
            print(f"{FAIL} Không mở được file: {e}")
        return False
    if verbose:
        report(filepath, profile, results)
    return not [r for r in results if r[0] == FAIL]


def report(filepath, profile, results) -> tuple[int, int, int]:
    print(f"=== Validate thể thức: {Path(filepath).name}  (profile: {profile}) ===\n")
    ok = warn = fail = 0
    for status, label, detail in results:
        print(f"  {status} {label}: {detail}")
        if status == OK:
            ok += 1
        elif status == WARN:
            warn += 1
        else:
            fail += 1
    print(f"\nTổng: {OK}{ok}  {WARN}{warn}  {FAIL}{fail}  / {len(results)} mục")
    if fail:
        print(f"{FAIL} CÓ LỖI NẶNG — KHÔNG bàn giao khi chưa sửa.")
    elif warn:
        print(f"{WARN} Không có lỗi nặng; mục ⚠ phải soi tay "
              "(references/validation-checklist.md mục B).")
    else:
        print(f"{OK} Sạch ở mức script kiểm được. Vẫn phải soi mắt bản render.")
    return ok, warn, fail


_STATUS_WORD = {OK: "OK", WARN: "WARN", FAIL: "FAIL"}


def results_to_json(filepath, profile, results) -> dict:
    """Dạng máy-đọc-được cho agent tự động sửa lỗi (self-healing) mà không cần parse chuỗi text.

    status tổng thể: "PASSED" (0 FAIL) | "PASSED_WITH_WARNINGS" (0 FAIL, có WARN) | "FAILED" (có FAIL).
    Mỗi item: {status, label, detail} — KHÔNG có field expected/actual/fix_suggestion máy-móc vì
    phần lớn lỗi thể thức (font, lề, placeholder...) không quy về 1 con số đơn giản để "tự sửa mù" —
    agent vẫn phải đọc `detail` (tiếng Việt, đã đủ rõ) rồi quyết định sửa gì.
    """
    items = [{"status": _STATUS_WORD[s], "label": label, "detail": detail}
             for s, label, detail in results]
    fail_n = sum(1 for i in items if i["status"] == "FAIL")
    warn_n = sum(1 for i in items if i["status"] == "WARN")
    overall = "FAILED" if fail_n else ("PASSED_WITH_WARNINGS" if warn_n else "PASSED")
    return {
        "file": str(filepath),
        "profile": profile,
        "overall_status": overall,
        "counts": {"ok": len(items) - fail_n - warn_n, "warn": warn_n, "fail": fail_n, "total": len(items)},
        "items": items,
    }


def main() -> int:
    ap = argparse.ArgumentParser(description="Validate .docx theo thể thức NĐ 30/2020/NĐ-CP")
    ap.add_argument("filepath")
    ap.add_argument("--profile", default="administrative",
                    help="administrative | bieu-mau-noi-bo | minutes-administrative "
                         "| academic | general")
    ap.add_argument("--allow-placeholder", action="store_true",
                    help="placeholder chỉ là cảnh báo — dùng khi validate MẪU trong templates/")
    ap.add_argument("--json", action="store_true",
                    help="In kết quả dạng JSON (machine-readable) ra stdout thay vì báo cáo chữ — "
                         "cho agent tự động đọc và quyết định sửa, không cần parse text tiếng Việt.")
    args = ap.parse_args()

    path = Path(args.filepath).resolve()
    if not path.is_file():
        if args.json:
            print(json.dumps({"overall_status": "ERROR", "error": f"không thấy file: {path}"},
                              ensure_ascii=False))
        else:
            print(f"ERROR: không thấy file: {path}", file=sys.stderr)
        return 2

    try:
        results = run_checks(path, args.profile, args.allow_placeholder)
    except Exception as e:
        if args.json:
            print(json.dumps({"overall_status": "ERROR", "error": f"Không mở được file: {e}"},
                              ensure_ascii=False))
        else:
            print(f"{FAIL} Không mở được file: {e}", file=sys.stderr)
        return 2

    if args.json:
        payload = results_to_json(path, args.profile, results)
        print(json.dumps(payload, ensure_ascii=False, indent=2))
        return 2 if payload["overall_status"] == "FAILED" else (1 if payload["counts"]["warn"] else 0)

    _, warn, fail = report(path, args.profile, results)
    if fail:
        return 2
    return 1 if warn else 0


if __name__ == "__main__":
    sys.exit(main())
