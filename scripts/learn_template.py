"""learn_template — đọc 1 file Word VBHC, extract spec + validate ND30.

Đọc file mẫu của cơ quan, phân tích
thể thức, kiểm tra so với NĐ 30, sinh:
  - <basename>.spec.json — machine-readable spec để builder dựng lại template tương đương
  - <basename>.report.md — báo cáo người đọc (chỗ nào chuẩn / chỗ nào sai ND30)

Usage:
    python learn_template.py <file.docx>
    python learn_template.py <file.docx> --out <prefix>
"""
from __future__ import annotations

import argparse
import json
import re
import sys, io
from pathlib import Path

if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

sys.path.insert(0, str(Path(__file__).parent.resolve()))
from docx import Document  # noqa: E402

import validate_docx as vt  # noqa: E402
import find_placeholders as fp  # noqa: E402
import normalize_template as nt  # noqa: E402
from rules_loader import load_rules  # noqa: E402


# ============================================================
# Loại VB detection (theo keyword tên loại in hoa hoặc V/v)
# ============================================================

LOAI_VB_PATTERNS = [
    ("phieu-ghi-y-kien", r"PHIẾU\s+(GHI\s+Ý\s+KIẾN|BIỂU\s+QUYẾT)"),
    ("bao-cao",          r"^BÁO CÁO\s*$|^BÁO CÁO\b"),
    ("to-trinh",         r"^TỜ TRÌNH\s*$|^TỜ TRÌNH\b"),
    ("quyet-dinh",       r"^QUYẾT ĐỊNH\s*$|^QUYẾT ĐỊNH\b"),
    ("ke-hoach",         r"^KẾ HOẠCH\s*$|^KẾ HOẠCH\b"),
    ("thong-bao",        r"^THÔNG BÁO\s*$|^THÔNG BÁO\b"),
    ("huong-dan",        r"^HƯỚNG DẪN\s*$|^HƯỚNG DẪN\b"),
    ("bien-ban",         r"BIÊN BẢN"),
    ("giay-moi",         r"GIẤY MỜI"),
    ("cong-van",         r"V/v|V/V"),  # CV cuối cùng — fallback nếu có V/v
]


def detect_loai_vb(text: str) -> str:
    for loai, pat in LOAI_VB_PATTERNS:
        if re.search(pat, text, re.MULTILINE):
            return loai
    return "unknown"


# ============================================================
# Page setup
# ============================================================

def extract_page_setup(doc) -> dict:
    section = doc.sections[0]
    return {
        "page_width_cm":  round(section.page_width.cm,  2) if section.page_width else None,
        "page_height_cm": round(section.page_height.cm, 2) if section.page_height else None,
        "margins": {
            "top_cm":    round(section.top_margin.cm,    2),
            "bottom_cm": round(section.bottom_margin.cm, 2),
            "left_cm":   round(section.left_margin.cm,   2),
            "right_cm":  round(section.right_margin.cm,  2),
        },
    }


# ============================================================
# Cơ quan chủ quản + ban hành (lấy từ table 0)
# ============================================================

def extract_org_info(doc) -> dict:
    if not doc.tables:
        return {"co_quan_chu_quan": None, "co_quan_ban_hanh": None}
    t = doc.tables[0]
    if not t.rows or not t.rows[0].cells:
        return {"co_quan_chu_quan": None, "co_quan_ban_hanh": None}
    left = t.rows[0].cells[0]
    lines = [p.text.strip() for p in left.paragraphs if p.text.strip()]
    return {
        "co_quan_chu_quan": lines[0] if lines else None,
        "co_quan_ban_hanh": lines[1] if len(lines) >= 2 else None,
    }


# ============================================================
# Số / ký hiệu / V/v
# ============================================================

def extract_so_kyhieu(text: str) -> dict:
    m = re.search(r"Số:\s*([0-9]*)\s*/\s*([A-ZĐa-z&\-_]+)", text)
    if not m:
        return {"so_vb": None, "ky_hieu": None}
    return {"so_vb": m.group(1).strip(), "ky_hieu": m.group(2).strip()}


def extract_vv(text: str):
    m = re.search(r"V/v\s+([^\n]+)", text)
    return m.group(1).strip() if m else None


# ============================================================
# Địa danh + ngày tháng năm
# ============================================================

DATE_RE = re.compile(
    r"([^\n,]+),\s*ngày\s*([\d]*)\s*tháng\s*([\d]*)\s*năm\s*([\d]*)"
)


def extract_dia_danh_ngay(text: str):
    m = DATE_RE.search(text)
    if not m:
        return None
    # Phần trước dấu "," có thể chứa text khác (vd "Hạnh phúc\nTuyên Quang") — lấy
    # token cuối sau ngắt dòng.
    raw = m.group(1).strip()
    dia_danh = raw.split("\n")[-1].strip()
    return {
        "dia_danh": dia_danh,
        "ngay": m.group(2).strip(),
        "thang": m.group(3).strip(),
        "nam": m.group(4).strip(),
    }


# ============================================================
# Người ký + Quyền hạn
# ============================================================

CHUC_VU_RE = re.compile(
    r"(KT\.|TL\.|TUQ\.)?\s*"
    r"(GIÁM ĐỐC|PHÓ GIÁM ĐỐC|CHỦ TỊCH|PHÓ CHỦ TỊCH|"
    r"CHÁNH VĂN PHÒNG|PHÓ CHÁNH VĂN PHÒNG|"
    r"TRƯỞNG PHÒNG|PHÓ TRƯỞNG PHÒNG|"
    r"BỘ TRƯỞNG|THỨ TRƯỞNG|"
    r"VỤ TRƯỞNG|CỤC TRƯỞNG|VIỆN TRƯỞNG)"
)


def extract_nguoi_ky(text: str):
    m = CHUC_VU_RE.search(text)
    if not m:
        return None
    return {"quyen_han": (m.group(1) or "").strip(), "chuc_vu": m.group(2).strip()}


# ============================================================
# Phòng soạn (từ Lưu) + Người soạn (trong ngoặc)
# ============================================================

def extract_phong_soan(text: str) -> dict:
    m = re.search(r"Lưu:\s*VT\s*,\s*([A-Za-zĐ&\-_]+)(?:\s*\(([^)]+)\))?", text)
    if not m:
        return {"phong_viet_tat": None, "nguoi_soan": None}
    return {
        "phong_viet_tat": m.group(1).strip(),
        "nguoi_soan": (m.group(2) or "").strip() or None,
    }


# ============================================================
# Nơi nhận (list các dòng "- ...")
# ============================================================

def extract_noi_nhan(text: str) -> list[str]:
    m = re.search(r"Nơi nhận:\s*\n((?:.*\n?)+?)(?=Lưu:|\Z)", text)
    if not m:
        return []
    block = m.group(1)
    items = []
    for ln in block.split("\n"):
        s = ln.strip()
        if s.startswith("-"):
            items.append(s.lstrip("-").strip())
    return items


# ============================================================
# Phán đoán điểm cần điều chỉnh so với ND30
# ============================================================

def assess_against_nd30(spec: dict, raw_text: str = "") -> list[dict]:
    """Trả về list issues: mỗi issue {level: 'ok'|'warn'|'fix', topic, message}.

    Mục đích: chỉ ra cho user chỗ nào nên điều chỉnh khi chuyển từ mẫu user
    sang template chuẩn (ví dụ: lề sai, font sai, typo, ...).
    """
    issues = []
    # Typo + encoding fixes — load từ tri-thuc-template/rules/typo-fixes.yaml
    # (fallback hardcode 2 rule cơ bản nếu YAML thiếu để giữ backward compat).
    typo_data = load_rules("typo-fixes") or {
        "encoding_fixes": [{
            "find": "Ð",
            "topic": "Ký tự Đ sai encoding",
            "message": "Phát hiện 'Ð' (U+00D0) — phải đổi sang 'Đ' (U+0110) khi chuẩn hóa",
            "level": "fix",
        }],
        "typo_fixes": [{
            "find": "kính giửi",
            "topic": "Lỗi chính tả",
            "message": "'kính giửi' → đúng là 'kính gửi'",
            "level": "fix",
        }],
    }
    for rule in (typo_data.get("encoding_fixes") or []) + (typo_data.get("typo_fixes") or []):
        if rule["find"] in raw_text:
            issues.append({
                "level": rule.get("level", "fix"),
                "topic": rule.get("topic", "Lỗi chính tả"),
                "message": rule.get("message", f"Phát hiện '{rule['find']}'"),
            })

    m = spec["page"]["margins"]
    if (m["left_cm"], m["right_cm"], m["top_cm"], m["bottom_cm"]) == (3, 2, 2, 2):
        issues.append({"level": "ok", "topic": "Lề",
                       "message": "Lề 3-2-2-2 khớp NĐ30"})
    else:
        issues.append({"level": "fix", "topic": "Lề",
                       "message": f"Lề hiện tại {m} — NĐ30 quy định trái 3, phải 2, trên 2, dưới 2 cm"})

    sk = spec["so_kyhieu"]
    if spec["loai_vb"] == "phieu-ghi-y-kien":
        if not sk["so_vb"] and not sk["ky_hieu"]:
            issues.append({"level": "ok", "topic": "Số VB",
                           "message": "Phiếu nội bộ — không cần Số VB ✓"})
    elif sk["ky_hieu"]:
        if spec["loai_vb"] == "cong-van":
            if "-" in sk["ky_hieu"] or sk["ky_hieu"].count("/") == 0:
                issues.append({"level": "ok", "topic": "Ký hiệu CV",
                               "message": f"Ký hiệu '{sk['ky_hieu']}' (CV không có chữ tắt loại — đúng)"})
        else:
            # VB khác CV phải có chữ tắt loại trong ký hiệu (BC, QĐ, TTr, ...)
            issues.append({"level": "ok", "topic": "Ký hiệu",
                           "message": f"Ký hiệu '{sk['ky_hieu']}'"})
    else:
        issues.append({"level": "warn", "topic": "Số VB", "message": "Không phát hiện Số/Ký hiệu"})

    nk = spec.get("nguoi_ky")
    if nk and nk["quyen_han"]:
        issues.append({"level": "ok", "topic": "Quyền hạn ký",
                       "message": f"Ký thay/thừa lệnh: {nk['quyen_han']} {nk['chuc_vu']}"})

    ps = spec["phong_soan"]
    if spec["loai_vb"] != "phieu-ghi-y-kien":
        if ps["phong_viet_tat"]:
            issues.append({"level": "ok", "topic": "Phòng soạn",
                           "message": f"Phòng soạn: {ps['phong_viet_tat']}"
                                      + (f" (người soạn: {ps['nguoi_soan']})" if ps["nguoi_soan"] else "")})
        else:
            issues.append({"level": "warn", "topic": "Phòng soạn",
                           "message": "Không phát hiện 'Lưu: VT, <phòng>.' — cần bổ sung"})

    return issues


# ============================================================
# Build human-readable Markdown report
# ============================================================

def build_report(spec: dict, validation: list, issues: list) -> str:
    L = []
    L.append(f"# Báo cáo học mẫu — {spec['file']}")
    L.append("")

    L.append(f"## Loại VB phát hiện: **{spec['loai_vb']}**")
    L.append("")

    L.append("## 1. Cơ quan + Số/Ký hiệu")
    org = spec["org"]
    L.append(f"- Cơ quan chủ quản: `{org['co_quan_chu_quan']}`")
    L.append(f"- Cơ quan ban hành: `{org['co_quan_ban_hanh']}`")
    sk = spec["so_kyhieu"]
    if sk["ky_hieu"]:
        L.append(f"- Số: `{sk['so_vb']}/{sk['ky_hieu']}`")
    if spec.get("vv"):
        L.append(f"- Trích yếu V/v: `{spec['vv'][:120]}...`" if len(spec['vv']) > 120
                 else f"- Trích yếu V/v: `{spec['vv']}`")
    L.append("")

    L.append("## 2. Địa danh + Ngày")
    dn = spec.get("dia_danh_ngay")
    if dn:
        L.append(f"- `{dn['dia_danh']}, ngày {dn['ngay'] or '___'} tháng {dn['thang']} năm {dn['nam']}`")
    L.append("")

    L.append("## 3. Người ký + Phòng soạn")
    nk = spec.get("nguoi_ky")
    if nk:
        if nk["quyen_han"]:
            L.append(f"- Quyền hạn ký: `{nk['quyen_han']}`")
        L.append(f"- Chức vụ: `{nk['chuc_vu']}`")
    ps = spec["phong_soan"]
    if ps["phong_viet_tat"]:
        L.append(f"- Phòng soạn: `{ps['phong_viet_tat']}`"
                 + (f" — người soạn: `{ps['nguoi_soan']}`" if ps["nguoi_soan"] else ""))
    L.append("")

    L.append("## 4. Nơi nhận")
    for it in spec.get("noi_nhan", []):
        L.append(f"- {it}")
    if not spec.get("noi_nhan"):
        L.append("- (không có / phiếu nội bộ)")
    L.append("")

    L.append("## 5. Page setup")
    p = spec["page"]
    L.append(f"- Khổ: {p.get('page_width_cm')}cm × {p.get('page_height_cm')}cm")
    m = p["margins"]
    L.append(f"- Lề: trái {m['left_cm']} | phải {m['right_cm']} | trên {m['top_cm']} | dưới {m['bottom_cm']} cm")
    L.append("")

    L.append("## 6. Validate ND30 (9 thành phần)")
    for label, (status, detail) in validation:
        L.append(f"- {status} **{label}** — {detail}")
    L.append("")

    L.append("## 7. Đề xuất điều chỉnh khi chuyển sang template chuẩn")
    icon = {"ok": "✓", "warn": "⚠", "fix": "✗"}
    for iss in issues:
        L.append(f"- {icon.get(iss['level'], '·')} **{iss['topic']}** — {iss['message']}")
    L.append("")

    # NEW v0.10: Placeholders detected
    L.append("## 8. Placeholders [KEY] phát hiện")
    placeholders = spec.get("placeholders", {}).get("placeholders", [])
    if placeholders:
        L.append(f"Tổng: **{len(placeholders)}** placeholders unique. Liệt kê:")
        for p in placeholders:
            L.append(f"- `{p['key']}` — {p['count']} lần (vd: {p['occurs'][0]['snippet']})")
    else:
        L.append("- (file không có placeholder dạng `[KEY]` — đây có thể là VB đã fill đầy đủ, không phải template)")
    L.append("")

    # NEW v0.10: Auto-fixes sẽ áp dụng khi save làm template
    L.append("## 9. Auto-fixes sẽ áp dụng (khi `vbhc_update_template`)")
    auto = spec.get("auto_fixable", {})
    if any(v > 0 for v in auto.values()):
        for k, v in sorted(auto.items(), key=lambda x: -x[1]):
            if v > 0:
                L.append(f"- {k}: **{v}** lần")
    else:
        L.append("- (không có lỗi cơ học — file đã sạch)")
    L.append("")

    # NEW v0.10: Manual review needed
    L.append("## 10. Cần user review thủ công (KHÔNG auto-fix)")
    manual = spec.get("manual_review_needed", [])
    if manual:
        for it in manual:
            L.append(f"- **{it['topic']}**: hiện tại `{it['current']}` — chuẩn `{it['expected']}`")
            L.append(f"  → {it['action']}")
    else:
        L.append("- (không có)")
    L.append("")

    return "\n".join(L)


# ============================================================
# Main
# ============================================================

def detect_manual_review(spec: dict) -> list[dict]:
    """Detect các điểm CẦN USER REVIEW THỦ CÔNG (không auto-fix được).
    Vd: lề khác chuẩn ND30 — có thể là phong cách cơ quan (VD lề 3.5cm cho VB
    nhiều dấu mộc), hoặc đơn vị đo khác.
    """
    items = []
    m = spec["page"]["margins"]
    if (m["left_cm"], m["right_cm"], m["top_cm"], m["bottom_cm"]) != (3, 2, 2, 2):
        items.append({
            "topic": "Lề trang",
            "current": f"trái {m['left_cm']} | phải {m['right_cm']} | trên {m['top_cm']} | dưới {m['bottom_cm']} cm",
            "expected": "trái 3 | phải 2 | trên 2 | dưới 2 cm (NĐ30)",
            "action": "User xác nhận có muốn đổi page setup không (giữ nguyên = phong cách cơ quan, sửa = chuẩn NĐ30)",
        })
    return items


def learn(path: Path):
    doc = Document(str(path))
    text = vt.collect_all_text(doc)

    spec = {
        "file": path.name,
        "loai_vb": detect_loai_vb(text),
        "page": extract_page_setup(doc),
        "org": extract_org_info(doc),
        "so_kyhieu": extract_so_kyhieu(text),
        "vv": extract_vv(text),
        "dia_danh_ngay": extract_dia_danh_ngay(text),
        "nguoi_ky": extract_nguoi_ky(text),
        "phong_soan": extract_phong_soan(text),
        "noi_nhan": extract_noi_nhan(text),
    }

    # NEW v0.10: Placeholders + auto-fixable issues + manual review
    spec["placeholders"] = fp.scan(path)
    spec["auto_fixable"] = nt.dry_run(path)
    spec["manual_review_needed"] = detect_manual_review(spec)

    validation = [
        ("1. Quốc hiệu + Tiêu ngữ",   vt.check_quoc_hieu(text)),
        ("2. Tên cơ quan ban hành",   vt.check_co_quan(text)),
        ("3. Số/ký hiệu",             vt.check_so_van_ban(text)),
        ("4. Tên loại + Trích yếu",   vt.check_ten_loai(text)),
        ("5. Nội dung",               vt.check_noi_dung(text)),
        ("6. Người ký",               vt.check_nguoi_ky(text)),
        ("7. Dấu/chữ ký số",          vt.check_dau()),
        ("8. Nơi nhận + Lưu",         vt.check_noi_nhan(text)),
        ("9. Phụ lục",                vt.check_phu_luc(text)),
    ]
    spec["validation_summary"] = [
        {"label": label, "status": status, "detail": detail}
        for label, (status, detail) in validation
    ]

    issues = assess_against_nd30(spec, raw_text=text)
    spec["assessment"] = issues
    return spec, validation, issues


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("file")
    ap.add_argument("--out", help="Output prefix (default: cùng folder, basename)")
    args = ap.parse_args()

    path = Path(args.file).resolve()
    if not path.is_file():
        print(f"ERROR: file not found: {path}", file=sys.stderr)
        return 1

    spec, validation, issues = learn(path)

    out_prefix = Path(args.out).resolve() if args.out else path.with_suffix("")
    spec_path = out_prefix.with_suffix(".spec.json")
    report_path = out_prefix.with_suffix(".report.md")

    spec_path.write_text(json.dumps(spec, ensure_ascii=False, indent=2), encoding="utf-8")
    report_path.write_text(build_report(spec, validation, issues), encoding="utf-8")

    print(f"[OK] Spec:   {spec_path}")
    print(f"[OK] Report: {report_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main() or 0)
