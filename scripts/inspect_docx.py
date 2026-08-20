"""Inspect a .docx file's structure: paragraphs + tables + cells.

Usage:
    python inspect_docx.py <file.docx>

Used for debugging template fill — find the exact paragraph index / table /
row / col that contains text you need to replace.
"""
import sys
from pathlib import Path

from _common import slugify_vn  # noqa: F401  # imports stdout fix
from docx import Document


def main():
    if len(sys.argv) < 2:
        print("Usage: python inspect_docx.py <file.docx>", file=sys.stderr)
        return 1

    path = Path(sys.argv[1]).resolve()
    if not path.is_file():
        print(f"ERROR: file not found: {path}", file=sys.stderr)
        return 1

    doc = Document(str(path))

    print(f"=== {path.name} ===\n")

    print("--- PARAGRAPHS ---")
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        runs_info = f" [{len(p.runs)} runs]" if len(p.runs) > 1 else ""
        if t:
            print(f"[P{i}]{runs_info} {t[:140]}")
        else:
            print(f"[P{i}] (empty)")

    print(f"\n--- TABLES ({len(doc.tables)} table(s)) ---")
    for ti, table in enumerate(doc.tables):
        print(f"\n>> Table {ti}: {len(table.rows)} rows × {len(table.columns)} cols")
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                txt = cell.text.replace("\n", " ⏎ ").strip()
                paragraphs_count = len(cell.paragraphs)
                runs_count = sum(len(p.runs) for p in cell.paragraphs)
                meta = f" [{paragraphs_count}p/{runs_count}r]"
                if txt:
                    print(f"  T{ti}/R{ri}/C{ci}{meta} {txt[:140]}")
                else:
                    print(f"  T{ti}/R{ri}/C{ci}{meta} (empty)")

    return 0


if __name__ == "__main__":
    sys.exit(main())
