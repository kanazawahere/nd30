"""Fill a .docx template with data — handles both plain paragraphs AND table cells.

Usage modes:

    # Mode 1: replace by exact match (works for paragraphs; falls back to cells)
    python fill_template.py <template> <output> --replace "OLD_TEXT" "NEW_TEXT" [--replace ...]

    # Mode 2: target specific cell by table/row/col coordinates
    python fill_template.py <template> <output> --cell T0 R1 C1 "NEW_TEXT" [--cell ...]

    # Mode 3: target specific paragraph by index
    python fill_template.py <template> <output> --para 7 "NEW_TEXT" [--para ...]

    # Mode 4: combine
    python fill_template.py <template> <output> \\
        --cell T1 R1 C1 "Dự thảo Nghị quyết..." \\
        --para 7 "- Ý kiến khác (nếu có): Nhất trí dự thảo." \\
        --replace "ngày 09 tháng 9 năm 2025" "ngày 09 tháng 5 năm 2026"

Why this script exists:
    Text trong ô bảng của .docx (nhất là file convert từ .doc) thường bị chẻ thành
    nhiều run — search-replace kiểu thường không khớp. Script này sửa thẳng từng run
    nên fill được cả ô bảng.

Notes:
    - Preserves formatting of the FIRST run; clears subsequent runs in the same paragraph.
    - When replacing a paragraph in a cell, removes any extra paragraphs after.
    - --replace iterates ALL paragraphs and ALL cells.
"""
import argparse
import shutil
import sys
from pathlib import Path

from _common import slugify_vn  # noqa: F401  # imports stdout fix
from docx import Document


def replace_paragraph_text(p, new_text):
    """Replace all text in a paragraph, keeping format of first run."""
    if not p.runs:
        p.add_run(new_text)
        return
    first = p.runs[0]
    first.text = new_text
    for r in p.runs[1:]:
        r.text = ""


def cell_set_text(cell, new_text):
    """Set cell text to new_text — collapse all paragraphs into the first one."""
    if not cell.paragraphs:
        cell.add_paragraph(new_text)
        return
    first_p = cell.paragraphs[0]
    replace_paragraph_text(first_p, new_text)
    # Remove extra paragraphs in the cell
    for p_extra in list(cell.paragraphs[1:]):
        p_extra._element.getparent().remove(p_extra._element)


def search_and_replace_doc(doc, find: str, repl: str) -> int:
    """Replace text in all paragraphs (top level) and all cells. Returns count of replacements."""
    count = 0
    # Top-level paragraphs
    for p in doc.paragraphs:
        full = "".join(r.text for r in p.runs)
        if find in full:
            new_full = full.replace(find, repl)
            replace_paragraph_text(p, new_full)
            count += 1
    # Cells in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full = cell.text
                if find in full:
                    cell_set_text(cell, full.replace(find, repl))
                    count += 1
    return count


def parse_coord(s):
    """Parse 'T0' / 'R1' / 'C2' format → int."""
    if len(s) < 2 or s[0] not in "TRC":
        raise argparse.ArgumentTypeError(f"Bad coord '{s}' (expected T0/R1/C2)")
    try:
        return int(s[1:])
    except ValueError:
        raise argparse.ArgumentTypeError(f"Bad coord '{s}'")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("template")
    ap.add_argument("output")
    ap.add_argument("--replace", nargs=2, metavar=("OLD", "NEW"), action="append", default=[],
                    help="Replace OLD with NEW in all paragraphs and cells")
    ap.add_argument("--cell", nargs=4, metavar=("Tx", "Rx", "Cx", "TEXT"), action="append", default=[],
                    help="Set text of cell at table/row/col coords")
    ap.add_argument("--para", nargs=2, metavar=("INDEX", "TEXT"), action="append", default=[],
                    help="Set text of top-level paragraph at given index")
    args = ap.parse_args()

    src = Path(args.template).resolve()
    dst = Path(args.output).resolve()

    if not src.is_file():
        print(f"ERROR: template not found: {src}", file=sys.stderr)
        return 1

    # Copy template → output (preserve original)
    if src != dst:
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(src, dst)

    doc = Document(str(dst))

    # Apply --cell ops
    for tspec, rspec, cspec, text in args.cell:
        ti = parse_coord(tspec)
        ri = parse_coord(rspec)
        ci = parse_coord(cspec)
        cell = doc.tables[ti].rows[ri].cells[ci]
        cell_set_text(cell, text)
        print(f"  ✓ Cell T{ti}/R{ri}/C{ci} ← {text[:60]}")

    # Apply --para ops
    for idx, text in args.para:
        i = int(idx)
        replace_paragraph_text(doc.paragraphs[i], text)
        print(f"  ✓ Paragraph P{i} ← {text[:60]}")

    # Apply --replace ops
    for old, new in args.replace:
        n = search_and_replace_doc(doc, old, new)
        if n:
            print(f"  ✓ Replaced {n} occurrence(s): {old[:40]} → {new[:40]}")
        else:
            print(f"  ⚠ Not found: {old[:60]}")

    doc.save(str(dst))
    print(f"\n✓ Saved: {dst}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
