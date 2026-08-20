"""build_docx — helper python-docx dựng văn bản hành chính đúng thể thức NĐ 30/2020/NĐ-CP.

nha-van:exempt — đây là mã nguồn, không phải văn bản gửi ra ngoài.

Nguồn: nền lấy từ `vbhc/scripts/vbhc_doc_builder.py` (biencuong, Unlicense/public domain);
phần "helper chung" ở cuối file do skill này viết thêm. Xem `NGUON-GOC.md`.

Module helper cho việc xây dựng .docx VBHC với XML chuẩn xác, đặc biệt:
  - Khối quốc hiệu/tên cơ quan với GẠCH CHÂN NGẮN bên dưới (tên CQ ban hành + Hạnh phúc)
  - Table không có border (header layout)
  - Column widths sticky (fix bug python-docx)
  - Line spacing chuẩn ND30
  - Font Times New Roman + East Asian font khai báo đúng

Why: python-docx default không tự gen XML đầy đủ cho các điểm trên.
"""
from __future__ import annotations

import sys, io
if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ============================================================
# Low-level XML helpers
# ============================================================

def _make_element(tag: str, attrs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attrs:
        for k, v in attrs.items():
            el.set(qn(k), str(v))
    return el


def set_paragraph_borders(p, *, top=None, bottom=None, left=None, right=None):
    """Add borders to a paragraph. Each side: dict {sz, val, color, space} or None.

    Default border values: val='single', sz=6 (3/4 pt), color='000000', space=1.
    """
    pPr = p._p.get_or_add_pPr()
    # Remove existing pBdr if any
    existing = pPr.find(qn('w:pBdr'))
    if existing is not None:
        pPr.remove(existing)
    pBdr = _make_element('w:pBdr')
    for side, conf in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        if conf is None:
            continue
        attrs = {
            'w:val': conf.get('val', 'single'),
            'w:sz': str(conf.get('sz', 6)),
            'w:space': str(conf.get('space', 1)),
            'w:color': conf.get('color', '000000'),
        }
        pBdr.append(_make_element(f'w:{side}', attrs))
    pPr.append(pBdr)


def set_paragraph_indent(p, *, left_twips: int | None = None, right_twips: int | None = None,
                         first_line_twips: int | None = None):
    """Set paragraph indentation in twips (1cm ≈ 567 twips, 1pt = 20 twips)."""
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = _make_element('w:ind')
        pPr.append(ind)
    if left_twips is not None:
        ind.set(qn('w:left'), str(left_twips))
    if right_twips is not None:
        ind.set(qn('w:right'), str(right_twips))
    if first_line_twips is not None:
        ind.set(qn('w:firstLine'), str(first_line_twips))


def set_paragraph_spacing(p, *, before_pt: float | None = None, after_pt: float | None = None,
                          line_pt: float | None = None, line_rule: str = 'auto'):
    """Set paragraph spacing. before/after in points; line spacing as multiplier (1.5) or pt with rule='exact'."""
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = _make_element('w:spacing')
        pPr.append(spacing)
    if before_pt is not None:
        spacing.set(qn('w:before'), str(int(before_pt * 20)))
    if after_pt is not None:
        spacing.set(qn('w:after'), str(int(after_pt * 20)))
    if line_pt is not None:
        if line_rule == 'auto':
            # In 'auto', line is in 240ths (single=240, 1.5=360, double=480)
            spacing.set(qn('w:line'), str(int(line_pt * 240)))
            spacing.set(qn('w:lineRule'), 'auto')
        else:
            spacing.set(qn('w:line'), str(int(line_pt * 20)))
            spacing.set(qn('w:lineRule'), line_rule)


def remove_table_borders(table):
    """Remove all borders from a table (4 sides + insideH + insideV)."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        return
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = _make_element('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tblBorders.append(_make_element(f'w:{side}', {'w:val': 'nil'}))
    tblPr.append(tblBorders)


def set_table_cell_margins_zero(table):
    """Set table-level cell margins to 0 to give cells full width.

    Word default cell margin is ~0.19cm left+right = 0.38cm total wasted per cell.
    For tight layouts (header table), set to 0 to use full column width.
    """
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        return
    existing = tblPr.find(qn('w:tblCellMar'))
    if existing is not None:
        tblPr.remove(existing)
    cellMar = _make_element('w:tblCellMar')
    for side in ('top', 'left', 'bottom', 'right'):
        cellMar.append(_make_element(f'w:{side}', {'w:w': '0', 'w:type': 'dxa'}))
    tblPr.append(cellMar)


def set_cell_borders_visible(cell, *, color='000000', sz=4):
    """Set 4-side single border on 1 cell (for góp ý table cells)."""
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = _make_element('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        tcBorders.append(_make_element(f'w:{side}', {
            'w:val': 'single', 'w:sz': str(sz), 'w:space': '0', 'w:color': color
        }))
    tcPr.append(tcBorders)


def align_table_to_left_margin(table):
    """Force table to align flush with the page's left margin (no indent).

    python-docx default tables sometimes inherit a small left indent (~0.19cm) from
    Normal style, making the table not flush with text. This sets tblInd to 0
    AND alignment LEFT, guaranteeing the table starts exactly at the left margin.
    """
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        return
    # Remove existing tblInd
    existing = tblPr.find(qn('w:tblInd'))
    if existing is not None:
        tblPr.remove(existing)
    # Add tblInd = 0
    tblInd = _make_element('w:tblInd', {'w:w': '0', 'w:type': 'dxa'})
    tblPr.append(tblInd)


def set_table_column_widths(table, widths_cm: list[float]):
    """Force column widths on EVERY cell (fix python-docx bug where setting columns[i].width
    only affects first cell, not entire column)."""
    # Set on tblGrid first
    tbl = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = _make_element('w:tblGrid')
        tbl.insert(0, tblGrid)
    # Clear existing gridCol
    for gc in tblGrid.findall(qn('w:gridCol')):
        tblGrid.remove(gc)
    for w_cm in widths_cm:
        gc = _make_element('w:gridCol', {'w:w': str(int(w_cm * 567))})
        tblGrid.append(gc)
    # Set on each cell
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci >= len(widths_cm):
                break
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = _make_element('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(int(widths_cm[ci] * 567)))
            tcW.set(qn('w:type'), 'dxa')
    # Also set table layout to fixed
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        layout = tblPr.find(qn('w:tblLayout'))
        if layout is None:
            layout = _make_element('w:tblLayout')
            tblPr.append(layout)
        layout.set(qn('w:type'), 'fixed')


# ============================================================
# Run / paragraph builders
# ============================================================

def add_run(p, text: str, *, bold=False, italic=False, underline=False,
            size_pt: float = 13, font: str = 'Times New Roman'):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if underline:
        r.underline = True
    r.font.size = Pt(size_pt)
    r.font.name = font
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = _make_element('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:cs'), font)
    rFonts.set(qn('w:eastAsia'), font)
    return r


# ============================================================
# Page setup (NĐ 30/2020)
# ============================================================

def setup_page(doc):
    """Lề chuẩn NĐ 30: trên/dưới/phải = 2cm, trái = 3cm. Khổ A4."""
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = _make_element('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')


# ============================================================
# Header block: cơ quan chủ quản/ban hành | quốc hiệu/tiêu ngữ
# ============================================================

def add_short_underline_paragraph(cell, *, indent_cm: float = 1.5,
                                   space_before_pt: float = 1.2,
                                   stroke_sz: int = 4):
    """Add an empty paragraph in cell with short top-border (gạch ngắn ở giữa).

    Paragraph border-top is a true vector line in Word — equivalent to inserting
    a thin horizontal line shape. Width controlled by indent on both sides.

    Args:
        space_before_pt: gap from text above to the line (1.2pt theo yêu cầu).
        stroke_sz: border weight in 1/8 pt (sz=4 → 0.5pt = 1/2pt).
    """
    p = cell.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_indent(p,
                        left_twips=int(indent_cm * 567),
                        right_twips=int(indent_cm * 567))
    set_paragraph_spacing(p, before_pt=space_before_pt, after_pt=2,
                          line_pt=1, line_rule='exact')
    set_paragraph_borders(p, top={'val': 'single', 'sz': stroke_sz,
                                   'color': '000000', 'space': 0})
    return p


def add_header_section(doc, *,
                       co_quan_chu_quan: str,
                       co_quan_ban_hanh: str,
                       quoc_hieu: str = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
                       tieu_ngu: str = "Độc lập - Tự do - Hạnh phúc",
                       co_quan_size_pt: float = 13,     # NĐ 30: 12-13pt
                       quoc_hieu_size_pt: float = 12,   # NĐ 30: 12-13pt (chọn 12 để 1 dòng)
                       tieu_ngu_size_pt: float = 14,    # NĐ 30: 13-14pt
                       left_col_cm: float = 7.0,
                       right_col_cm: float = 9.0,
                       cq_underline_pct: float = 0.55,    # 55% width tên đơn vị
                       qh_underline_pct: float = 1.00):   # 100% width tiêu ngữ
    """Add the standard top header: 1 row × 2 cols, no borders, with short underlines.

    Default widths: 7cm (left) + 9cm (right) = 16cm — fits A4 with NĐ 30 margins
    (page=21cm, margins 3+2=5cm, content area 16cm). Cell padding = 0.

    Default font sizes per NĐ 30/2020 Phụ lục I:
      - Tên cơ quan: 12-13pt (default 13pt)
      - Quốc hiệu: 12-13pt (default 12pt to fit 1 line in 9cm cell)
      - Tiêu ngữ: 13-14pt (default 14pt)
    """
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    set_table_column_widths(table, [left_col_cm, right_col_cm])
    remove_table_borders(table)
    set_table_cell_margins_zero(table)
    align_table_to_left_margin(table)

    # ----- Cell trái -----
    left_cell = table.rows[0].cells[0]
    # Clear default empty paragraph
    left_cell.text = ""
    p1 = left_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Line spacing single, before/after = 0 (theo yêu cầu user)
    set_paragraph_spacing(p1, before_pt=0, after_pt=0, line_pt=1.0, line_rule='auto')
    add_run(p1, co_quan_chu_quan, bold=False, size_pt=co_quan_size_pt)

    p2 = left_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p2, before_pt=0, after_pt=0, line_pt=1.0, line_rule='auto')
    add_run(p2, co_quan_ban_hanh, bold=True, size_pt=co_quan_size_pt)

    # Gạch chân dưới tên cơ quan ban hành: chiếm cq_underline_pct của TEXT WIDTH
    # (estimated). Times New Roman 13pt bold uppercase: ~0.27cm per char.
    cq_text_width = len(co_quan_ban_hanh) * 0.27
    cq_underline_width = cq_text_width * cq_underline_pct
    # Center-align: each side indent
    cq_indent_each_side = (left_col_cm - cq_underline_width) / 2
    cq_indent_each_side = max(cq_indent_each_side, 0.3)  # min 0.3cm gutter
    add_short_underline_paragraph(left_cell, indent_cm=cq_indent_each_side,
                                   space_before_pt=1.2, stroke_sz=4)

    # ----- Cell phải -----
    right_cell = table.rows[0].cells[1]
    right_cell.text = ""
    p1 = right_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p1, before_pt=0, after_pt=0, line_pt=1.0, line_rule='auto')
    add_run(p1, quoc_hieu, bold=True, size_pt=quoc_hieu_size_pt)

    p2 = right_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p2, before_pt=0, after_pt=0, line_pt=1.0, line_rule='auto')
    add_run(p2, tieu_ngu, bold=True, size_pt=tieu_ngu_size_pt)

    # Gạch chân dưới tiêu ngữ: chiếm qh_underline_pct của TEXT WIDTH "Độc lập - Tự do - Hạnh phúc"
    # Times New Roman 14pt bold: ~0.20cm per char (text "Độc lập - Tự do - Hạnh phúc" 26 chars ~ 5.2cm)
    qh_text_width = len(tieu_ngu) * 0.20
    qh_underline_width = qh_text_width * qh_underline_pct
    qh_indent_each_side = (right_col_cm - qh_underline_width) / 2
    qh_indent_each_side = max(qh_indent_each_side, 0.3)
    add_short_underline_paragraph(right_cell, indent_cm=qh_indent_each_side,
                                   space_before_pt=1.2, stroke_sz=4)

    return table


def add_centered_title_with_underline(doc, text: str, *, size_pt: float = 14,
                                       underline_indent_cm: float = 4.5):
    """Standalone centered title with short underline (used for biểu mẫu — phiếu).

    Vd: "PHIẾU GHI Ý KIẾN THÀNH VIÊN UBND TỈNH" + gạch chân ngắn dưới.
    Gap giữa text và gạch chân: 2.5pt.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_pt=12, after_pt=0, line_pt=1.0, line_rule='auto')
    add_run(p, text, bold=True, size_pt=size_pt)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_indent(p2, left_twips=int(underline_indent_cm * 567),
                         right_twips=int(underline_indent_cm * 567))
    # gap 1.2pt + stroke 0.5pt (sz=4)
    set_paragraph_spacing(p2, before_pt=1.2, after_pt=8, line_pt=1, line_rule='exact')
    set_paragraph_borders(p2, top={'val': 'single', 'sz': 4, 'color': '000000', 'space': 0})


# ============================================================
# Page numbering — số trang trên header, ẩn trang 1, bắt đầu trang 2 = số 2
# ============================================================

def apply_page_numbering(doc, *, hide_first_page: bool = True,
                         start_at: int = 1, position: str = 'top_center'):
    """Bật page numbers trong header. Ẩn trang đầu nếu hide_first_page=True.

    Args:
        hide_first_page: True → trang 1 không hiện số trang
        start_at: số bắt đầu (default 1 → trang 2 hiện "2")
        position: 'top_center' (chỉ option supported hiện tại)

    Note: Word "Different First Page" phải bật trên section để hide_first_page work.
    """
    section = doc.sections[0]
    section.different_first_page_header_footer = hide_first_page

    # Set startingPageNumber via sectPr/pgNumType
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = _make_element('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:start'), str(start_at))

    # First page header: empty (hide page number on page 1)
    if hide_first_page:
        first_header = section.first_page_header
        # Just leave it empty — Word's "Different First Page" handles hiding
        if not first_header.paragraphs:
            first_header.add_paragraph()

    # Default header: insert PAGE field, centered
    default_header = section.header
    # Clear existing content
    if default_header.paragraphs:
        for p in default_header.paragraphs[1:]:
            p._element.getparent().remove(p._element)
        p = default_header.paragraphs[0]
        # Clear runs in first paragraph
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
    else:
        p = default_header.add_paragraph()

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Insert PAGE field
    run = p.add_run()
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'

    # Build XML: <w:fldChar w:fldCharType="begin"/> <w:instrText>PAGE</w:instrText> <w:fldChar w:fldCharType="end"/>
    fldChar_begin = _make_element('w:fldChar', {'w:fldCharType': 'begin'})
    instrText = _make_element('w:instrText')
    instrText.text = 'PAGE   \\* MERGEFORMAT'
    instrText.set(qn('xml:space'), 'preserve')
    fldChar_end = _make_element('w:fldChar', {'w:fldCharType': 'end'})

    run._element.append(fldChar_begin)
    run._element.append(instrText)
    run._element.append(fldChar_end)


# ============================================================
# Số văn bản + Địa danh, ngày
# ============================================================

def add_so_vb_and_date_section(doc, *,
                                so_vb: str = "",
                                ky_hieu: str = "",  # vd: "BC-SGDĐT"
                                trich_yeu: str = "",  # cho công văn (V/v ...)
                                dia_danh: str = "",
                                ngay: int | str = "",
                                thang: int | str = "",
                                nam: int | str = "",
                                is_cong_van: bool = False,
                                left_col_cm: float = 7.0,
                                right_col_cm: float = 9.0):
    """Add the Số/V/v + Địa danh ngày block as a 2-col table (no borders).

    For Công văn: Số:.../X-Y kèm V/v ... bên trái.
    For others: Số:.../X-Y bên trái (không có V/v).
    Ngày tháng năm bên phải, italic.
    """
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    set_table_column_widths(table, [left_col_cm, right_col_cm])
    remove_table_borders(table)
    set_table_cell_margins_zero(table)
    align_table_to_left_margin(table)

    # Cell trái: Số + (V/v if cong van)
    left_cell = table.rows[0].cells[0]
    left_cell.text = ""
    p1 = left_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p1, before_pt=6, after_pt=0, line_pt=1.15, line_rule='auto')
    if so_vb:
        so_text = f"Số: {so_vb}/{ky_hieu}"
    else:
        so_text = f"Số:        /{ky_hieu}"
    add_run(p1, so_text, size_pt=13)

    if is_cong_van and trich_yeu:
        p2 = left_cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p2, before_pt=0, after_pt=0, line_pt=1.15, line_rule='auto')
        # NĐ30: trích yếu V/v in THƯỜNG ĐỨNG, không nghiêng (chỉ địa danh-ngày italic)
        add_run(p2, f"V/v {trich_yeu}", italic=False, size_pt=13)

    # Cell phải: Địa danh, ngày tháng năm
    # Rule (NĐ 30 + chính sách user): KHÔNG điền trước NGÀY ban hành (do VPHC điền khi ban hành).
    # Nhưng tháng/năm có thể điền theo current để tránh "tháng     năm   " quá trống.
    right_cell = table.rows[0].cells[1]
    right_cell.text = ""
    p1 = right_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p1, before_pt=6, after_pt=0, line_pt=1.15, line_rule='auto')

    now = datetime.now()
    # Tháng/năm: dùng giá trị user truyền, nếu không truyền → fill current
    thang_val = thang if thang else now.month
    nam_val = nam if nam else now.year
    thang_str = f"{int(thang_val):02d}" if isinstance(thang_val, int) or (isinstance(thang_val, str) and str(thang_val).isdigit()) else str(thang_val)
    nam_str = str(nam_val)

    if ngay:
        # User truyền ngày cụ thể → điền hết
        ngay_str = f"{int(ngay):02d}" if isinstance(ngay, int) or (isinstance(ngay, str) and str(ngay).isdigit()) else str(ngay)
        date_text = f"{dia_danh}, ngày {ngay_str} tháng {thang_str} năm {nam_str}"
    else:
        # Bỏ trống ngày, điền tháng + năm
        date_text = f"{dia_danh}, ngày        tháng {thang_str} năm {nam_str}"
    add_run(p1, date_text, italic=True, size_pt=13)

    return table


# ============================================================
# Title block (tên loại + trích yếu)
# ============================================================

def add_title_block(doc, *, ten_loai: str, trich_yeu: str = "",
                    underline_pct: float = 0.35,  # gạch chân dưới trích yếu: 30-40%
                    content_width_cm: float = 16.0):
    """Tên loại VB IN HOA + trích yếu (cho VB không phải Công văn).

    Theo yêu cầu user:
    - Spacing trong khổ trích yếu = 0 (single, before/after = 0)
    - Có gạch chân ngắn dưới trích yếu, chiếm 30-40% width dòng trên
    """
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # before=18 để cách section trên, after=0 (tight với trích yếu)
    set_paragraph_spacing(p1, before_pt=18, after_pt=0, line_pt=1.0, line_rule='auto')
    add_run(p1, ten_loai.upper(), bold=True, size_pt=14)

    if trich_yeu:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # spacing = 0 trong khổ này (theo yêu cầu)
        set_paragraph_spacing(p2, before_pt=0, after_pt=0, line_pt=1.0, line_rule='auto')
        add_run(p2, trich_yeu, bold=True, size_pt=14)

        # Gạch chân ngắn dưới trích yếu: 30-40% width
        # Gap 1.2pt + stroke 0.5pt (sz=4)
        underline_width = content_width_cm * underline_pct
        indent_each_side = (content_width_cm - underline_width) / 2
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_indent(p3,
                             left_twips=int(indent_each_side * 567),
                             right_twips=int(indent_each_side * 567))
        set_paragraph_spacing(p3, before_pt=1.2, after_pt=8, line_pt=1, line_rule='exact')
        set_paragraph_borders(p3,
                              top={'val': 'single', 'sz': 4, 'color': '000000', 'space': 0})


# ============================================================
# Kính gửi
# ============================================================

def add_kinh_gui(doc, recipient: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_pt=6, after_pt=6, line_pt=1.5, line_rule='auto')
    add_run(p, f"Kính gửi: {recipient}", size_pt=13)


# ============================================================
# Body paragraph (nội dung thường)
# ============================================================

def add_body_paragraph(doc, text: str, *, indent_first_cm: float = 1.1,
                      align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_pt: float = 1.0,
                      space_before_pt: float = 6, space_after_pt: float = 6,
                      apply_widow_control: bool = True, italic: bool = False,
                      bold: bool = False, size_pt: float = 13):
    """Body paragraph chuẩn ND 30:
    - Justify (căn 2 đầu)
    - Indent đầu dòng 1.1cm
    - Spacing before=after=6pt → line spacing PHẢI = single (1.0) để không dãn quá
    - Auto compress char spacing -0.1pt nếu có nguy cơ widow (2 từ mồ côi)

    italic: dùng cho dòng "Nghị quyết này đã được HĐND ... thông qua ngày..." của VBQPPL
    (Phụ lục I NĐ 78/2025, Mục III.2.c: chữ nghiêng, đặt phía dưới điều cuối cùng).
    """
    p = doc.add_paragraph()
    p.alignment = align
    set_paragraph_spacing(p, before_pt=space_before_pt, after_pt=space_after_pt,
                          line_pt=line_pt, line_rule='auto')
    if indent_first_cm:
        set_paragraph_indent(p, first_line_twips=int(indent_first_cm * 567))
    add_run(p, text, size_pt=size_pt, italic=italic, bold=bold)
    if apply_widow_control:
        _apply_widow_compress(p, text)
    return p


def add_section_heading(doc, text: str, *, level: int = 1, italic: bool | None = None,
                        bold: bool | None = None,
                        indent_first_cm: float = 1.1):
    """Section heading: tất cả các cấp đều IN ĐẬM, KHÔNG NGHIÊNG (theo NĐ30).

    Args:
        level: 1 / 2 / 3 — chỉ để phân loại logic, không thay đổi style mặc định
        bold/italic: override nếu cần (vd cố ý nghiêng cho 1 trường hợp đặc biệt)
    """
    default_bold, default_italic = True, False
    if bold is None:
        bold = default_bold
    if italic is None:
        italic = default_italic

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before_pt=6, after_pt=2, line_pt=1.0, line_rule='auto')
    if indent_first_cm:
        set_paragraph_indent(p, first_line_twips=int(indent_first_cm * 567))
    add_run(p, text, bold=bold, italic=italic, size_pt=13)
    return p


def _apply_widow_compress(p, text: str, threshold_chars: int = 12):
    """Tránh "từ mồ côi" (orphan words) — câu cuối khổ chỉ còn 1-2 từ ngắn.

    Heuristic: ước lượng số ký tự dòng cuối = len(text) % 75. Nếu trong khoảng
    (0, threshold_chars*2] = (0, 24] → khả năng dòng cuối có 1-2 từ mồ côi rất
    cao → nén char spacing -0.1pt cho TOÀN BỘ paragraph để dồn từ về dòng trên.

    Threshold mặc định 12 (= 24 chars dòng cuối — tương đương 4-5 từ tiếng Việt
    có dấu) đủ rộng để bắt phần lớn trường hợp thực tế.
    """
    if len(text) < 80:
        return
    chars_per_line_approx = 75  # A4 lề 3-2 + indent 1.1 + cỡ 13pt ≈ 75 char/line
    last_line_chars = len(text) % chars_per_line_approx
    if 0 < last_line_chars <= threshold_chars * 2:
        # Char spacing condensed -0.1pt = -2 twentieths of pt
        for r in p.runs:
            rPr = r._element.get_or_add_rPr()
            spacing = rPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = _make_element('w:spacing')
                rPr.append(spacing)
            spacing.set(qn('w:val'), '-2')


# ============================================================
# Signature + Nơi nhận block (2-col table, no borders)
# ============================================================

def add_signature_noi_nhan(doc, *,
                            noi_nhan_items: list[str],
                            chuc_vu: str,
                            nguoi_ky: str,
                            quyen_han: str = "",   # "" / "KT." / "TL." / "TUQ." / "TM."
                            chuc_vu_thay: str = "",  # vd: "GIÁM ĐỐC" khi quyen_han='KT.'
                            phong_viet_tat: str = "",  # vd: "GDPT" → "Lưu: VT, GDPT."
                            empty_lines_for_signature: int = 5):
    """Khối cuối: Nơi nhận trái + chữ ký phải.

    Nguyên tắc theo yêu cầu user:
    - KT.GĐ/PGĐ: 2 dòng chức vụ — line spacing single, before/after = 0 (tight)
    - Nơi nhận: nếu phong_viet_tat → tự thêm "Lưu: VT, <phong_viet_tat>." cuối list
    - Spacing-after sau dòng cuối Nơi nhận = 1 dòng

    quyen_han + chuc_vu_thay: nếu KT.GIÁM ĐỐC + PHÓ GIÁM ĐỐC ký:
        quyen_han='KT.', chuc_vu_thay='GIÁM ĐỐC', chuc_vu='PHÓ GIÁM ĐỐC'

    "TM." (thay mặt tập thể, vd Quyết định của UBND — Phụ lục I NĐ78/2025 Mẫu 19/20):
        quyen_han='TM.', chuc_vu_thay='ỦY BAN NHÂN DÂN', chuc_vu='CHỦ TỊCH'
        → render "TM. ỦY BAN NHÂN DÂN" dòng 1, "CHỦ TỊCH" dòng 2. Khác "KT." về ngữ nghĩa
        (KT. = ký thay 1 người cụ thể vắng mặt; TM. = thay mặt cả tập thể cơ quan) nhưng
        cùng cấu trúc trình bày 2 dòng nên dùng chung tham số.
    """
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    set_table_column_widths(table, [8.0, 8.0])
    remove_table_borders(table)
    # nd30: cấm chẻ khối "Nơi nhận | chữ ký" ngang qua 2 trang (đã thấy lỗi khi render)
    _trPr = table.rows[0]._tr.get_or_add_trPr()
    _trPr.append(_make_element('w:cantSplit'))

    # ----- Cột trái: Nơi nhận -----
    left = table.rows[0].cells[0]
    left.text = ""
    p1 = left.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p1, before_pt=12, after_pt=0, line_pt=1.15, line_rule='auto')
    add_run(p1, "Nơi nhận:", bold=True, italic=True, size_pt=11)

    # Auto-add "Lưu: VT, <viet_tat>." nếu user cung cấp phong_viet_tat
    items_final = list(noi_nhan_items)
    # Sanitize: NĐ 30 KHÔNG quy định "(để báo cáo)"/"(để phối hợp)" trong dấu ngoặc.
    # Nếu user (hoặc AI) lỡ truyền vào → tự strip phần trong "(...)" cuối câu.
    import re as _re
    # nd30: chỉ bỏ phần trong ngoặc, GIỮ dấu câu cuối (NĐ30 yêu cầu mục cuối kết bằng '.')
    items_final = [_re.sub(r"\s*\([^)]*\)\s*(?=[.;]?\s*$)", "", it).rstrip() for it in items_final]
    items_final = [it if it.endswith((";", ".")) else it + ";" for it in items_final]
    if items_final:
        items_final[-1] = items_final[-1].rstrip(";") + ("" if items_final[-1].endswith(".") else ".")
    if phong_viet_tat and not any("Lưu" in i and "VT" in i for i in items_final):
        items_final.append(f"- Lưu: VT, {phong_viet_tat}.")

    for idx, item in enumerate(items_final):
        p = left.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # spacing-after = 12pt (1 line) cho dòng cuối, 0 cho các dòng khác
        is_last = (idx == len(items_final) - 1)
        after = 12 if is_last else 0
        set_paragraph_spacing(p, before_pt=0, after_pt=after, line_pt=1.15, line_rule='auto')
        add_run(p, item, size_pt=11)

    # ----- Cột phải: chức vụ + người ký (TIGHT, spacing 0) -----
    right = table.rows[0].cells[1]
    right.text = ""

    if quyen_han:
        # 2 dòng chức vụ — line spacing single, before/after = 0
        p1 = right.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p1, before_pt=12, after_pt=0, line_pt=1.0, line_rule='auto')
        add_run(p1, f"{quyen_han} {chuc_vu_thay}".strip(), bold=True, size_pt=13)

        p2 = right.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p2, before_pt=0, after_pt=0, line_pt=1.0, line_rule='auto')
        add_run(p2, chuc_vu, bold=True, size_pt=13)
    else:
        p1 = right.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p1, before_pt=12, after_pt=0, line_pt=1.0, line_rule='auto')
        add_run(p1, chuc_vu, bold=True, size_pt=13)

    # Empty lines for signature (tight spacing)
    for _ in range(empty_lines_for_signature):
        p = right.add_paragraph()
        set_paragraph_spacing(p, before_pt=0, after_pt=0, line_pt=1.15, line_rule='auto')

    # Tên người ký (tight spacing)
    p = right.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_pt=0, after_pt=0, line_pt=1.0, line_rule='auto')
    add_run(p, nguoi_ky, bold=True, size_pt=13)


# ============================================================
# Bảng góp ý (4 cột với border, dùng cho báo cáo góp ý dự thảo)
# ============================================================

def _auto_resize_columns(rows, n_cols, total_width_cm=16.0,
                         min_width_cm=2.8, max_width_cm=6.0,
                         headers=None):
    """Auto-resize column widths based on content length.

    Iterative algorithm:
    1. Compute weight per column from content (avg*0.6 + max*0.4) + header length
    2. Distribute total_width by weight
    3. If any column < min OR > max → clamp it, remove from pool, redistribute
       remaining width among unclamped columns by their weights
    4. Repeat until no violations.

    This guarantees: sum(widths) == total_width AND all widths in [min, max].
    """
    if not rows or n_cols == 0:
        return [total_width_cm / max(n_cols, 1)] * n_cols

    # 1. Weights — header length carries higher multiplier (1.8) to prevent header wrap
    weights = []
    for ci in range(n_cols):
        lens = [len(str(row[ci])) for row in rows if ci < len(row)]
        if not lens:
            lens = [10]
        avg = sum(lens) / len(lens)
        mx = max(lens)
        w = avg * 0.6 + mx * 0.4
        if headers and ci < len(headers):
            # Header weight: chars * 1.8 — header should never wrap
            w = max(w, len(str(headers[ci])) * 1.8)
        weights.append(max(w, 1.0))

    # 2. Iterative clamping
    fixed: dict[int, float] = {}
    available = total_width_cm

    while True:
        free_cols = [i for i in range(n_cols) if i not in fixed]
        if not free_cols:
            break
        free_weight = sum(weights[i] for i in free_cols)
        if free_weight <= 0:
            equal = available / len(free_cols)
            for i in free_cols:
                fixed[i] = equal
            break

        # Compute proposed widths for free cols
        proposed = {i: available * weights[i] / free_weight for i in free_cols}

        # Find smallest violator (most deviant from range)
        violator = None
        violator_value = None
        for i, w in proposed.items():
            if w < min_width_cm:
                if violator is None or (min_width_cm - w) > (min_width_cm - (violator_value or w)):
                    violator = i
                    violator_value = min_width_cm
            elif w > max_width_cm:
                if violator is None or (w - max_width_cm) > ((violator_value or w) - max_width_cm):
                    violator = i
                    violator_value = max_width_cm

        if violator is None:
            # All proposed widths within [min, max] — accept
            for i, w in proposed.items():
                fixed[i] = w
            break
        else:
            fixed[violator] = violator_value
            available -= violator_value

    return [fixed[i] for i in range(n_cols)]


def add_gop_y_table(doc, rows: list[tuple[str, str, str, str]],
                     headers: tuple[str, str, str, str] = (
                         "Điều, khoản", "Nội dung Dự thảo",
                         "Đề nghị sửa thành", "Căn cứ, lí do góp ý"),
                     col_widths_cm: tuple[float, float, float, float] | None = None,
                     auto_resize: bool = True):
    """Add a 4-column góp ý table with borders, header row in bold.

    Args:
        rows: list of (col1, col2, col3, col4) tuples — góp ý items
        headers: header row labels
        col_widths_cm: explicit column widths in cm. If None and auto_resize=True,
                       compute widths from content.
        auto_resize: if True (default), compute widths from content length.
    """
    if col_widths_cm is None:
        if auto_resize:
            col_widths_cm = _auto_resize_columns(rows, n_cols=4,
                                                  headers=list(headers))
        else:
            col_widths_cm = (3.0, 4.5, 4.5, 4.0)
    col_widths_cm = list(col_widths_cm)
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.autofit = False
    set_table_column_widths(table, list(col_widths_cm))
    align_table_to_left_margin(table)

    # Header row
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before_pt=2, after_pt=2, line_pt=1.15, line_rule='auto')
        add_run(p, h, bold=True, size_pt=12)
        set_cell_borders_visible(cell)

    # Data rows
    # - Cột 0 (Điều, khoản): CENTER bold (label ngắn, đẹp khi center)
    # - Cột 1-3 (Nội dung dự thảo / Đề nghị / Căn cứ): JUSTIFY (text dài)
    for ri, row_data in enumerate(rows, start=1):
        for ci, content in enumerate(row_data):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP if ci > 0 else WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            if ci == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, before_pt=2, after_pt=2, line_pt=1.2, line_rule='auto')
            add_run(p, content, bold=(ci == 0), size_pt=12)
            set_cell_borders_visible(cell)

    return table


# ============================================================
# Bảng biểu quyết (4 cột STT/Nội dung/Đồng ý/Không đồng ý)
# ============================================================

def add_bieu_quyet_table(doc, items: list[dict]):
    """Add the standard biểu quyết table for Phiếu ghi ý kiến.

    Args:
        items: list of dicts {stt, noi_dung, dong_y (bool), khong_dong_y (bool)}
    """
    table = doc.add_table(rows=1 + len(items), cols=4)
    table.autofit = False
    set_table_column_widths(table, [1.5, 9.5, 2.5, 2.5])
    align_table_to_left_margin(table)

    headers = ["STT", "Nội dung", "Đồng ý", "Không đồng ý"]
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before_pt=2, after_pt=2, line_pt=1.15, line_rule='auto')
        add_run(p, h, bold=True, size_pt=12)
        set_cell_borders_visible(cell)

    for ri, item in enumerate(items, start=1):
        # STT
        cell = table.rows[ri].cells[0]
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before_pt=2, after_pt=2, line_pt=1.15, line_rule='auto')
        add_run(p, str(item.get("stt", ri)), size_pt=12)
        set_cell_borders_visible(cell)
        # Nội dung
        cell = table.rows[ri].cells[1]
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(p, before_pt=2, after_pt=2, line_pt=1.2, line_rule='auto')
        add_run(p, item.get("noi_dung", ""), size_pt=12)
        set_cell_borders_visible(cell)
        # Đồng ý
        cell = table.rows[ri].cells[2]
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before_pt=2, after_pt=2, line_pt=1.15, line_rule='auto')
        if item.get("dong_y"):
            add_run(p, "X", bold=True, size_pt=14)
        set_cell_borders_visible(cell)
        # Không đồng ý
        cell = table.rows[ri].cells[3]
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before_pt=2, after_pt=2, line_pt=1.15, line_rule='auto')
        if item.get("khong_dong_y"):
            add_run(p, "X", bold=True, size_pt=14)
        set_cell_borders_visible(cell)

    return table


# ============================================================
# Convenience: Standard skeleton for a VBHC document
# ============================================================

def create_vbhc_skeleton(*,
                          co_quan_chu_quan: str,
                          co_quan_ban_hanh: str,
                          so_vb: str = "",
                          ky_hieu: str,
                          trich_yeu: str,
                          dia_danh: str,
                          ngay: int | str = "",
                          thang: int | str = "",
                          nam: int | str = "",
                          is_cong_van: bool = False,
                          ten_loai_in_hoa: str = "",  # bỏ trống nếu là CV
                          ):
    """Create a Document with standard VBHC header — ready for body insertion.

    Returns (doc, body_anchor_paragraph) — body should be inserted before
    the anchor (or anchor can be deleted and body appended).
    """
    doc = Document()
    setup_page(doc)
    add_header_section(doc,
                       co_quan_chu_quan=co_quan_chu_quan,
                       co_quan_ban_hanh=co_quan_ban_hanh)
    add_so_vb_and_date_section(doc,
                                so_vb=so_vb,
                                ky_hieu=ky_hieu,
                                trich_yeu=trich_yeu,
                                dia_danh=dia_danh,
                                ngay=ngay, thang=thang, nam=nam,
                                is_cong_van=is_cong_van)
    if not is_cong_van and ten_loai_in_hoa:
        add_title_block(doc, ten_loai=ten_loai_in_hoa, trich_yeu=trich_yeu)
    return doc


# ============================================================
# Helper chung (viết thêm cho skill nd30)
# ============================================================

PLACEHOLDER = "[CẦN BỔ SUNG: {}]"


def placeholder(what: str) -> str:
    """Chuỗi placeholder chuẩn — validate_docx.py sẽ bắt được và báo còn thiếu."""
    return PLACEHOLDER.format(what)


def setup_document(doc=None, profile: str = "administrative"):
    """Tạo/cấu hình Document theo profile. Trả về doc.

    profile hành chính → lề 3/2/2/2 cm, A4, Times New Roman 13pt đen, thân bài canh đều.
    Các profile khác giữ nguyên hình thức trang nhưng không ràng buộc canh đều.
    """
    if doc is None:
        doc = Document()
    setup_page(doc)  # A4 + lề NĐ30 + font Normal (hàm của phần nền)

    style = doc.styles["Normal"]
    style.font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    if profile in ("administrative", "bieu-mau-noi-bo", "minutes-administrative"):
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)          # NĐ30 PL I II.6.e: tối thiểu 6pt
    pf.line_spacing = 1.0

    # Heading dùng chung font/màu — tránh Word tự bôi xanh
    for name, size in (("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 13)):
        if name in [s.name for s in doc.styles]:
            h = doc.styles[name]
            h.font.name = "Times New Roman"
            h.font.size = Pt(size)
            h.font.bold = True
            h.font.color.rgb = RGBColor(0, 0, 0)
            h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return doc


def add_two_col_block(doc, left_lines, right_lines, *,
                      left_col_cm: float = 8.0, right_col_cm: float = 8.0,
                      left_align=WD_ALIGN_PARAGRAPH.LEFT,
                      right_align=WD_ALIGN_PARAGRAPH.CENTER,
                      align_left_margin: bool = True):
    """Bảng 2 cột ẨN VIỀN — cách duy nhất đáng tin để xếp 2 khối cạnh nhau trong .docx.

    python-docx không có "tab stop 2 cột" tự nhiên; dùng dấu cách/tab tay sẽ lệch khi
    đổi máy hoặc đổi font. Đây là helper tổng quát cho mọi khối 2 cột:
      - đầu văn bản: tên cơ quan (trái) | Quốc hiệu + Tiêu ngữ (phải)
      - cuối văn bản: Nơi nhận (trái) | chức vụ + chữ ký (phải)

    Mỗi phần tử của left_lines/right_lines là str, hoặc dict:
        {"text": ..., "bold": bool, "italic": bool, "size_pt": float,
         "align": WD_ALIGN_PARAGRAPH.*}

    Khối chuyên dụng đã dựng sẵn: `add_header_section()` và `add_signature_noi_nhan()`.
    """
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    set_table_column_widths(table, [left_col_cm, right_col_cm])
    remove_table_borders(table)
    set_table_cell_margins_zero(table)
    if align_left_margin:
        align_table_to_left_margin(table)

    for cell, lines, default_align in (
        (table.rows[0].cells[0], left_lines, left_align),
        (table.rows[0].cells[1], right_lines, right_align),
    ):
        cell.text = ""
        first = True
        for item in lines:
            spec = {"text": item} if isinstance(item, str) else dict(item)
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            p.alignment = spec.get("align", default_align)
            set_paragraph_spacing(p, before_pt=0, after_pt=0, line_pt=1.0, line_rule="auto")
            add_run(p, spec.get("text", ""),
                    bold=spec.get("bold", False),
                    italic=spec.get("italic", False),
                    size_pt=spec.get("size_pt", 13))
    return table


def add_can_cu(doc, items: list[str]):
    """Các dòng 'Căn cứ ...' — in NGHIÊNG (NĐ30 PL I II.6.a).

    Tự đặt dấu: các dòng trên `;`, dòng cuối `.` — lỗi kinh điển ở cơ quan.
    """
    out = []
    n = len(items)
    for i, raw in enumerate(items):
        text = raw.rstrip(" ;.")
        text += "." if i == n - 1 else ";"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(p, before_pt=0, after_pt=6, line_pt=1.0, line_rule="auto")
        set_paragraph_indent(p, first_line_twips=int(1.1 * 567))
        add_run(p, text, italic=True, size_pt=13)
        out.append(p)
    return out


def add_bullet(doc, text: str, level: int = 1, profile: str = "administrative"):
    """Gạch đầu dòng bằng TIỀN TỐ TAY — NĐ30 không dùng bullet tự động của Word.

    level 1 → '- ', level 2 → '+ ', level 3 → '* ' (kèm thụt lề tương ứng).
    """
    prefix = {1: "- ", 2: "+ ", 3: "* "}.get(level, "- ")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, before_pt=0, after_pt=6, line_pt=1.0, line_rule="auto")
    if level > 1:
        set_paragraph_indent(p, left_twips=int((level - 1) * 1.1 * 567))
    add_run(p, prefix + text, size_pt=13)
    return p


def add_bullets(doc, items, level: int = 1, profile: str = "administrative"):
    return [add_bullet(doc, it, level, profile) for it in items]


def add_centered(doc, text: str, *, size_pt: float = 13, bold: bool = False,
                 italic: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_pt=6, after_pt=6, line_pt=1.0, line_rule="auto")
    add_run(p, text, bold=bold, italic=italic, size_pt=size_pt)
    return p


def add_table(doc, headers: list[str], rows: list[list], *,
              widths_cm: list[float] | None = None,
              profile: str = "administrative"):
    """Bảng có viền, hàng đầu in đậm, KHÔNG tô nền (profile hành chính cấm shading)."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.autofit = False
    if widths_cm is None:
        widths_cm = [16.0 / len(headers)] * len(headers)
    set_table_column_widths(tbl, widths_cm)
    for row in tbl.rows:
        for cell in row.cells:
            set_cell_borders_visible(cell)
    for ci, h in enumerate(headers):
        cell = tbl.cell(0, ci)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, str(h), bold=True, size_pt=13)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.cell(ri + 1, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, str(val), size_pt=13)
    return tbl


def add_ket_thuc(doc, text: str = "./."):
    """Dấu kết thúc văn bản hành chính."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, before_pt=0, after_pt=6, line_pt=1.0, line_rule="auto")
    add_run(p, text, size_pt=13)
    return p
