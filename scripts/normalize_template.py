"""normalize_template.py — auto-fix các lỗi cơ học khi import file .docx user
làm template chuẩn ND30.

Fixes thực hiện (an toàn, không đổi style/structure):
  - Đ encoding: U+00D0 ('Ð') → U+0110 ('Đ')  — common bug từ file convert .doc
  - Typo phổ biến: load từ tri-thuc-template/05-thong-tin-co-quan.yaml
    field `quy_tac_nd30.loi_chinh_ta_pho_bien`
  - Strip trailing whitespace trong cells (dọn rác visual)

KHÔNG fix:
  - Lề (cần user xác nhận thay đổi page setup)
  - Font/cỡ chữ (style cơ quan có thể khác mặc định)
  - Heading style (italic/bold) — phong cách cơ quan
  - Cấu trúc tables/paragraphs

Usage:
    python normalize_template.py <input.docx> <output.docx>
    python normalize_template.py <input.docx> --dry-run        # chỉ liệt kê

Module API:
    from normalize_template import normalize, dry_run
    counts = normalize(src_path, dst_path)             # dict {fix_name: count}
    counts = dry_run(src_path)                          # không write file
"""
from __future__ import annotations

import argparse
import shutil
import sys, io
from pathlib import Path

if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document


SKILL_DIR = Path(__file__).resolve().parent.parent
DEFAULT_RULES_YAML = SKILL_DIR / "tri-thuc-template" / "05-thong-tin-co-quan.yaml"


def _load_typo_rules(rules_yaml: Path) -> dict[str, str]:
    """Đọc map typo phổ biến từ yaml. Fallback hardcoded nếu file không có."""
    fallback = {"kính giửi": "kính gửi"}
    if not rules_yaml.is_file():
        return fallback
    try:
        import yaml
        data = yaml.safe_load(rules_yaml.read_text(encoding="utf-8")) or {}
    except Exception:
        return fallback
    rules = (data.get("quy_tac_nd30") or {}).get("loi_chinh_ta_pho_bien") or {}
    if not isinstance(rules, dict):
        return fallback
    return {**fallback, **{str(k): str(v) for k, v in rules.items()}}


# ============================================================
# Core scan/replace
# ============================================================

def _iter_runs(doc):
    """Yield (run, location_label) cho mọi run trong doc (paragraphs + table cells)."""
    for pi, p in enumerate(doc.paragraphs):
        for ri, r in enumerate(p.runs):
            yield r, f"paragraph_{pi}/run_{ri}"
    for ti, t in enumerate(doc.tables):
        for row_i, row in enumerate(t.rows):
            for col_i, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    for ri, r in enumerate(p.runs):
                        yield r, f"table_{ti}/r_{row_i}/c_{col_i}/p_{pi}/run_{ri}"


def _count_str(text: str, find: str) -> int:
    return text.count(find) if find else 0


def scan(doc, typo_rules: dict[str, str] | None = None) -> dict:
    """Scan toàn doc, đếm số lần mỗi fix sẽ áp dụng. KHÔNG sửa gì."""
    if typo_rules is None:
        typo_rules = _load_typo_rules(DEFAULT_RULES_YAML)

    counts = {"D_encoding": 0, "trailing_ws": 0}
    for typo in typo_rules:
        counts[f"typo:{typo}"] = 0

    for run, _loc in _iter_runs(doc):
        text = run.text or ""
        counts["D_encoding"] += _count_str(text, "Ð")
        for typo in typo_rules:
            counts[f"typo:{typo}"] += _count_str(text, typo)
        if text != text.rstrip() and text.strip():
            # has trailing whitespace and not just whitespace-only
            counts["trailing_ws"] += 1

    return counts


def apply_fixes(doc, typo_rules: dict[str, str] | None = None) -> dict:
    """Apply fixes IN-PLACE on doc. Trả về counts đã sửa."""
    if typo_rules is None:
        typo_rules = _load_typo_rules(DEFAULT_RULES_YAML)

    counts = {"D_encoding": 0, "trailing_ws": 0}
    for typo in typo_rules:
        counts[f"typo:{typo}"] = 0

    for run, _loc in _iter_runs(doc):
        original = run.text or ""
        new = original

        # Fix 1: Đ encoding
        n_d = new.count("Ð")
        if n_d:
            new = new.replace("Ð", "Đ")
            counts["D_encoding"] += n_d

        # Fix 2: typo replacements
        for typo, correct in typo_rules.items():
            n = new.count(typo)
            if n:
                new = new.replace(typo, correct)
                counts[f"typo:{typo}"] += n

        # Fix 3: trailing whitespace (only if non-empty meaningful content)
        stripped = new.rstrip()
        if new != stripped and stripped:
            new = stripped
            counts["trailing_ws"] += 1

        if new != original:
            run.text = new

    return counts


# ============================================================
# Public API
# ============================================================

def normalize(src: Path, dst: Path, typo_rules: dict[str, str] | None = None) -> dict:
    """Copy src → dst, apply fixes, save. Trả về counts."""
    src = Path(src)
    dst = Path(dst)
    dst.parent.mkdir(parents=True, exist_ok=True)
    if src.resolve() != dst.resolve():
        shutil.copyfile(src, dst)
    doc = Document(str(dst))
    counts = apply_fixes(doc, typo_rules)
    if any(v > 0 for v in counts.values()):
        doc.save(str(dst))
    return counts


def dry_run(src: Path, typo_rules: dict[str, str] | None = None) -> dict:
    """Scan src, return counts. KHÔNG write file."""
    doc = Document(str(src))
    return scan(doc, typo_rules)


# ============================================================
# CLI
# ============================================================

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input")
    ap.add_argument("output", nargs="?", help="Đích (nếu không pass + không dry-run, ghi đè input)")
    ap.add_argument("--dry-run", action="store_true", help="Chỉ scan, không write")
    args = ap.parse_args()

    src = Path(args.input).resolve()
    if not src.is_file():
        print(f"ERROR: file not found: {src}", file=sys.stderr)
        return 1

    rules = _load_typo_rules(DEFAULT_RULES_YAML)

    if args.dry_run:
        counts = dry_run(src, rules)
        print(f"=== Dry-run: {src.name} ===")
    else:
        dst = Path(args.output).resolve() if args.output else src
        counts = normalize(src, dst, rules)
        print(f"=== Normalized: {src.name} → {dst.name} ===")

    total = sum(counts.values())
    if total == 0:
        print("  (không có fix nào áp dụng — file đã sạch)")
    else:
        for name, n in sorted(counts.items(), key=lambda x: -x[1]):
            if n > 0:
                print(f"  {name}: {n} fix(es)")
        print(f"  TOTAL: {total} fix(es)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
